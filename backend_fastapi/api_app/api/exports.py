"""FastAPI 路由 - 导出功能 (迁移版)
支持 Word/Excel/PDF 格式导出
"""

from io import BytesIO
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse

from app.models.quotation import Quotation
from app.models.module import Module
from app.models.material import ModuleMaterial, Material
from app.models.fee import OtherFee, FeeType
from app.models.fee_rate import FeeRate
from app.models.exchange_rate import ExchangeRate
from app.models.user import User
from app.models.labor_hour import LaborHour
from app.models.travel_entry import PackingEntry, TravelPersonDays, TravelPersonTrip
from app.models.travel import TravelPersonTripFee
from api_app.main import get_db

# Word 导出
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Excel 导出
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# PDF 导出 - 使用 fpdf2 支持中文
from fpdf import FPDF

# 中文字体路径
import os
FONT_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'fonts')
FONT_REGULAR = os.path.join(FONT_DIR, 'simhei.ttf')
FONT_BOLD = os.path.join(FONT_DIR, 'simhei.ttf')

# 中英文标签映射
I18N = {
    'zh': {
        'quote': '报价单', 'basic_info': '基本信息', 'project_name': '项目名称',
        'scheme_no': '方案编号', 'business_owner': '业务负责人', 'created_at': '创建时间',
        'currency': '报价币种', 'material_list': '物料及费用清单',
        'module_col': '模块', 'item_col': '项目', 'spec_col': '规格',
        'category_col': '分类/品牌', 'unit_price_col': '单价', 'qty_col': '数量',
        'subtotal_col': '小计', 'total_col': '合计',
        'quote_summary': '报价汇总', 'project_col': '项目', 'amount_col': '金额',
        'material_total': '物料合计', 'fee_total': '费用合计', 'labor_total': '人力合计',
        'profit_rate': '利润率', 'subtotal_with_profit': '含利润率小计',
        'tax_rate': '税率', 'tax_amount': '税额', 'grand_total': '最终报价',
        'other_fees': '其他费用', 'unnamed_module': '未命名模块',
        'internal': '厂内', 'external': '厂外',
        'no_materials': '（无物料）', 'no_data': '无物料及费用数据',
        'labor_hours': '人力工时', 'version': '版本', 'version_no': '版本号',
        'profit_rate_label': '利润率', 'subtotal_label': '合计',
        'brand_col': '品牌', 'division_col': '设计工时划分', 'fee_name_col': '费用名称',
        'table1_total': '设备硬件合计', 'table2_total': '设备人力合计', 'table3_total': '其他合计',
        'profit_amount': '项目利润',
        'unit_col': '单位', 'unit_h': 'h', 'tax_included': '含税报价',
    },
    'en': {
        'quote': 'QUOTATION', 'basic_info': 'Basic Information', 'project_name': 'Project Name',
        'scheme_no': 'Scheme No.', 'business_owner': 'Sales Owner', 'created_at': 'Created At',
        'currency': 'Currency', 'material_list': 'Material & Cost List',
        'module_col': 'Module', 'item_col': 'Item', 'spec_col': 'Spec',
        'category_col': 'Category/Brand', 'unit_price_col': 'Unit Price', 'qty_col': 'Qty',
        'subtotal_col': 'Subtotal', 'total_col': 'Total',
        'quote_summary': 'Quote Summary', 'project_col': 'Item', 'amount_col': 'Amount',
        'material_total': 'Material Total', 'fee_total': 'Fee Total', 'labor_total': 'Labor Total',
        'profit_rate': 'Profit Rate', 'subtotal_with_profit': 'Subtotal w/ Profit',
        'tax_rate': 'Tax Rate', 'tax_amount': 'Tax Amount', 'grand_total': 'Grand Total',
        'other_fees': 'Other Fees', 'unnamed_module': 'Unnamed Module',
        'internal': 'Internal', 'external': 'External',
        'no_materials': '(No Materials)', 'no_data': 'No material or cost data',
        'labor_hours': 'Labor Hours', 'version': 'Version', 'version_no': 'Version No.',
        'profit_rate_label': 'Profit Rate', 'subtotal_label': 'Subtotal w/ Profit',
        'table1_total': 'Hardware Total', 'table2_total': 'Labor Total', 'table3_total': 'Other Total',
        'profit_amount': 'Project Profit',
        'unit_col': 'Unit', 'unit_h': 'h', 'tax_included': 'Tax Included',
    }
}

router = APIRouter(prefix='/api')


def t(key, lang='zh'):
    return I18N.get(lang, I18N['zh']).get(key, key)


def get_name(obj, lang='zh'):
    """获取对象的本地化名称，英文时优先用 name_en"""
    if lang == 'en':
        en = getattr(obj, 'name_en', None)
        if en:
            return en
    return obj.name


def get_fee_type_name(fee_type_id, lang='zh'):
    """根据 fee_type ID 或名称字符串获取费用类型名称"""
    if not fee_type_id:
        return fee_type_id or ''
    if isinstance(fee_type_id, int):
        from app import db as _flask_db
        ft = _flask_db.session.get(FeeType, fee_type_id)
        if ft:
            return ft.name_en if (lang == 'en' and ft.name_en) else ft.name
        return str(fee_type_id)
    ft = FeeType.query.filter(FeeType.name == str(fee_type_id)).first()
    if ft:
        return ft.name_en if (lang == 'en' and ft.name_en) else ft.name
    if lang == 'en':
        ft = FeeType.query.filter(FeeType.name_en == str(fee_type_id)).first()
        if ft:
            return ft.name_en
    return str(fee_type_id)


def get_currency_symbol(currency):
    symbols = {'CNY': '¥', 'USD': '$', 'EUR': '€', 'GBP': '£', 'JPY': '¥'}
    return symbols.get(currency, '¥')


def convert_currency(amount, to_currency, exchange_rates):
    if to_currency == 'CNY':
        return amount
    rate = exchange_rates.get(to_currency, 1)
    return amount / rate


def calculate_totals_with_rates(quotation, modules, fees, labor_hours):
    """计算报价单汇总（含费用系数）"""
    from app.models.fee_rate import FeeRate
    # 获取费用系数
    coeff = {}
    if hasattr(quotation, 'coefficients') and quotation.coefficients:
        coeff = quotation.coefficients
    else:
        rates = FeeRate.query.all()
        coeff = {fr.category: float(fr.rate) for fr in rates}
    rate_large = float(coeff.get('large', 1.0))
    rate_standard = float(coeff.get('standard', 1.0))
    rate_other = float(coeff.get('other', 1.0))

    # 物料合计（含系数）
    material_total = 0.0
    for module in modules:
        if hasattr(module, 'materials') and module.materials:
            for mm in module.materials:
                if not mm.material:
                    continue
                if mm.is_other:
                    unit_amount = float(mm.unit_price_override or float(mm.material.unit_price or 0))
                else:
                    rate = (rate_large if mm.material.category == 'large'
                            else (rate_standard if mm.material.category == 'standard' else rate_other))
                    unit_amount = float(mm.material.unit_price or 0) * rate * mm.quantity
                material_total += unit_amount

    # 费用合计
    fees_total = sum(float(f.amount or 0) for f in fees)

    # 人力合计
    labor_total = sum(float(l.total or 0) for l in labor_hours)

    # 利润
    profit_rate = float(quotation.profit_rate) if quotation.profit_rate else 0.0
    subtotal_before_profit = material_total + fees_total + labor_total
    profit_amount = subtotal_before_profit * profit_rate
    subtotal_with_profit = subtotal_before_profit + profit_amount

    # 税额
    tax_rate = float(quotation.tax_rate) if quotation.tax_rate else 0.0
    tax_amount = subtotal_with_profit * tax_rate

    grand_total = subtotal_with_profit + tax_amount

    return {
        'material_total_with_rates': round(material_total, 2),
        'fees_total': round(fees_total, 2),
        'labor_total': round(labor_total, 2),
        'profit_rate': profit_rate,
        'profit_amount': round(profit_amount, 2),
        'subtotal_with_profit': round(subtotal_with_profit, 2),
        'tax_rate': tax_rate,
        'tax_amount': round(tax_amount, 2),
        'grand_total': round(grand_total, 2),
        'fee_rates': coeff,
    }


# ==================== 导出 Word ====================

@router.get('/quotations/{quotation_id}/export/word')
def export_word(
    quotation_id: int,
    lang: str = Query('zh', description='语言: zh/en'),
    db=Depends(get_db),
):
    """导出报价单 Word 格式"""
    quotation = db.query(Quotation).get(quotation_id)
    if not quotation:
        raise HTTPException(status_code=404, detail='Quotation not found')

    # 线体：聚合所有子报价单数据
    if quotation.type == 'line':
        child_ids = [c.id for c in quotation.children.all()]
        all_ids = [quotation_id] + child_ids
        modules = Module.query.filter(Module.quotation_id.in_(all_ids)).all()
        for module in modules:
            module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()
        fees = OtherFee.query.filter(OtherFee.quotation_id.in_(all_ids)).all()
        labor_hours = LaborHour.query.filter(LaborHour.quotation_id.in_(all_ids)).all()
    else:
        modules = Module.query.filter_by(quotation_id=quotation_id).all()
        for module in modules:
            module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()
        fees = OtherFee.query.filter_by(quotation_id=quotation_id).all()
        labor_hours = LaborHour.query.filter_by(quotation_id=quotation_id).all()

    totals = calculate_totals_with_rates(quotation, modules, fees, labor_hours)

    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)

    doc = Document()

    # 标题
    title = doc.add_heading('报价单', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    doc.add_heading('基本信息', level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        (t('project_name', lang), quotation.name or ''),
        (t('scheme_no', lang), quotation.scheme_no or ''),
        (t('business_owner', lang), quotation.business_owner.real_name if quotation.business_owner else ''),
        ('创建时间', quotation.created_at.strftime('%Y-%m-%d %H:%M') if quotation.created_at else ''),
        (t('currency', lang), currency)
    ]
    for i, (key, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # 物料及费用清单
    doc.add_heading('物料及费用清单', level=1)

    module_groups = []
    for module in modules:
        items = []
        module_total = 0
        if module.materials:
            for mm in module.materials:
                mat = mm.material
                if mat:
                    item_total = float(mat.unit_price or 0) * mm.quantity
                    module_total += item_total
                    items.append({
                        'item': mat.name or '',
                        'spec': mat.spec or '',
                        'category': mat.brand or '',
                        'unit_price': f'¥{float(mat.unit_price or 0):.2f}',
                        'quantity': str(mm.quantity),
                        'subtotal': f'¥{item_total:.2f}',
                    })
        module_groups.append({
            'name': get_name(module, lang),
            'items': items,
            'total': module_total
        })

    fees_total = sum(float(f.amount or 0) for f in fees)
    labor_total_for_group = sum(float(l.total or 0) for l in labor_hours)
    fee_items = []
    for fee in fees:
        location_text = '厂内' if fee.location == 'internal' else ('厂外' if fee.location == 'external' else fee.location or '')
        fee_items.append({
            'item': get_fee_type_name(fee.fee_type, lang) or '',
            'spec': location_text,
            'category': get_fee_type_name(fee.fee_type, lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f'¥{float(fee.amount or 0):.2f}',
        })
    for l in labor_hours:
        fee_items.append({
            'item': l.name or '人力工时',
            'spec': f'{l.hours}h',
            'category': '人力工时',
            'unit_price': f'¥{float(l.unit_price or 0):.2f}/h',
            'quantity': '1',
            'subtotal': f'¥{float(l.total or 0):.2f}',
        })
    if fee_items:
        module_groups.append({
            'name': '其他费用',
            'items': fee_items,
            'total': fees_total + labor_total_for_group
        })

    if module_groups:
        combined_table = doc.add_table(rows=1, cols=8)
        combined_table.style = 'Table Grid'
        headers = ['模块', '项目', '规格', '分类/品牌', '单价', '数量', '小计', '合计']
        for j, header in enumerate(headers):
            combined_table.rows[0].cells[j].text = header

        for group in module_groups:
            if not group['items']:
                row = combined_table.add_row()
                row.cells[0].text = group['name']
                row.cells[1].text = '（无物料）'
                continue

            first_row = combined_table.add_row()
            first_row.cells[0].text = group['name']
            first_row.cells[1].text = group['items'][0]['item']
            first_row.cells[2].text = group['items'][0]['spec']
            first_row.cells[3].text = group['items'][0]['category']
            first_row.cells[4].text = group['items'][0]['unit_price']
            first_row.cells[5].text = group['items'][0]['quantity']
            first_row.cells[6].text = group['items'][0]['subtotal']
            first_row.cells[7].text = f'¥{group["total"]:.2f}'

            if len(group['items']) > 1:
                module_cell = first_row.cells[0]
                for i in range(1, len(group['items'])):
                    next_row = combined_table.add_row()
                    next_row.cells[1].text = group['items'][i]['item']
                    next_row.cells[2].text = group['items'][i]['spec']
                    next_row.cells[3].text = group['items'][i]['category']
                    next_row.cells[4].text = group['items'][i]['unit_price']
                    next_row.cells[5].text = group['items'][i]['quantity']
                    next_row.cells[6].text = group['items'][i]['subtotal']
                    module_cell = module_cell.merge(next_row.cells[0])

                total_cell = first_row.cells[7]
                for i in range(1, len(group['items'])):
                    next_row_idx = len(combined_table.rows) - 1
                    total_cell = total_cell.merge(combined_table.rows[next_row_idx].cells[7])
    else:
        doc.add_paragraph('无物料及费用数据')

    doc.add_paragraph()

    # 汇总信息
    doc.add_heading(t('quote_summary', lang), level=1)
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Table Grid'
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary_data = [
        (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
        (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
        (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
        (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
        (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
        (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
        (f'{t("grand_total", lang)}({currency})', f'{currency_symbol}{grand_total_converted:.2f}')
    ]
    for i, (key, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = value

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f'quotation_{quotation_id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'
    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


# ==================== 导出 Excel ====================

@router.get('/quotations/{quotation_id}/export/excel')
def export_excel(
    quotation_id: int,
    lang: str = Query('zh', description='语言: zh/en'),
    db=Depends(get_db),
):
    """导出报价单 Excel 格式"""
    quotation = db.query(Quotation).get(quotation_id)
    if not quotation:
        raise HTTPException(status_code=404, detail='Quotation not found')

    # 线体：聚合所有子报价单数据
    if quotation.type == 'line':
        child_ids = [c.id for c in quotation.children.all()]
        all_ids = [quotation_id] + child_ids
        modules = Module.query.filter(Module.quotation_id.in_(all_ids)).all()
        for module in modules:
            module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()
        fees = OtherFee.query.filter(OtherFee.quotation_id.in_(all_ids)).all()
        labor_hours = LaborHour.query.filter(LaborHour.quotation_id.in_(all_ids)).all()
    else:
        modules = Module.query.filter_by(quotation_id=quotation_id).all()
        for module in modules:
            module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()
        fees = OtherFee.query.filter_by(quotation_id=quotation_id).all()
        labor_hours = LaborHour.query.filter_by(quotation_id=quotation_id).all()

    totals = calculate_totals_with_rates(quotation, modules, fees, labor_hours)

    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)

    # 构建 module_groups
    module_groups = []
    for module in modules:
        items = []
        module_total = 0
        if module.materials:
            for mm in module.materials:
                mat = mm.material
                if mat:
                    item_total = float(mat.unit_price or 0) * mm.quantity
                    module_total += item_total
                    items.append({
                        'item': mat.name or '',
                        'spec': mat.spec or '',
                        'category': mat.brand or '',
                        'unit_price': f'¥{float(mat.unit_price or 0):.2f}',
                        'quantity': str(mm.quantity),
                        'subtotal': f'¥{item_total:.2f}',
                    })
        module_groups.append({
            'name': get_name(module, lang),
            'items': items,
            'total': module_total
        })

    fees_total = sum(float(f.amount or 0) for f in fees)
    labor_total_for_group = sum(float(l.total or 0) for l in labor_hours)
    fee_items = []
    for fee in fees:
        location_text = '厂内' if fee.location == 'internal' else ('厂外' if fee.location == 'external' else fee.location or '')
        fee_items.append({
            'item': get_fee_type_name(fee.fee_type, lang) or '',
            'spec': location_text,
            'category': get_fee_type_name(fee.fee_type, lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f'¥{float(fee.amount or 0):.2f}',
        })
    for l in labor_hours:
        fee_items.append({
            'item': l.name or '人力工时',
            'spec': f'{l.hours}h',
            'category': '人力工时',
            'unit_price': f'{float(l.unit_price or 0):.2f}',
            'quantity': '1',
            'subtotal': f'{float(l.total or 0):.2f}',
        })
    if fee_items:
        module_groups.append({
            'name': '其他费用',
            'items': fee_items,
            'total': fees_total + labor_total_for_group
        })

    wb = Workbook()
    wb.remove(wb.active)

    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Sheet 1: 基本信息
    ws1 = wb.create_sheet('基本信息')
    ws1.append(['报价单基本信息'])
    ws1.merge_cells('A1:D1')
    ws1['A1'].font = Font(bold=True, size=14)
    ws1['A1'].alignment = Alignment(horizontal='center')

    ws1.append([])
    basic_info = [
        (t('project_name', lang), quotation.name or ''),
        (t('scheme_no', lang), quotation.scheme_no or ''),
        (t('business_owner', lang), quotation.business_owner.real_name if quotation.business_owner else ''),
        ('创建时间', quotation.created_at.strftime('%Y-%m-%d %H:%M') if quotation.created_at else ''),
        (t('currency', lang), currency)
    ]
    for key, value in basic_info:
        ws1.append([key, value])

    ws1.append([])
    ws1.append(['费用系数'])
    ws1.merge_cells(f'A{ws1.max_row}:B{ws1.max_row}')
    ws1[f'A{ws1.max_row}'].font = Font(bold=True)

    rate_info = [
        ('大件系数', f"{totals['fee_rates'].get('large', 1.0)}x"),
        ('普通件系数', f"{totals['fee_rates'].get('standard', 1.0)}x"),
        ('其他件系数', f"{totals['fee_rates'].get('other', 1.0)}x")
    ]
    for key, value in rate_info:
        ws1.append([key, value])

    ws1.append([])
    ws1.append(['报价汇总'])
    ws1.merge_cells(f'A{ws1.max_row}:B{ws1.max_row}')
    ws1[f'A{ws1.max_row}'].font = Font(bold=True)

    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary = [
        ('物料合计', f'¥{totals["material_total_with_rates"]:.2f}'),
        ('费用合计', f'¥{fees_with_labor:.2f}'),
        ('利润率', f'{totals["profit_rate"] * 100:.1f}%'),
        ('含利润率小计', f'¥{totals["subtotal_with_profit"]:.2f}'),
        ('税率', f'{totals["tax_rate"] * 100:.0f}%'),
        ('税额', f'¥{totals["tax_amount"]:.2f}'),
        (f'最终报价({currency})', f'{currency_symbol}{grand_total_converted:.2f}')
    ]
    for key, value in summary:
        ws1.append([key, value])

    ws1.column_dimensions['A'].width = 18
    ws1.column_dimensions['B'].width = 30

    # Sheet 2: 物料及费用清单
    ws2 = wb.create_sheet('物料及费用清单')

    headers = ['模块', '项目', '规格', '分类/品牌', '单价', '数量', '小计', '合计']
    ws2.append(headers)
    for cell in ws2[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    row_num = 2
    merge_info = []

    for group in module_groups:
        if not group['items']:
            ws2.append([group['name'], '（无物料）', '', '', '', '', '', ''])
            for cell in ws2[row_num]:
                cell.border = thin_border
            row_num += 1
            continue

        start_row = row_num
        for i, item in enumerate(group['items']):
            ws2.append([
                group['name'] if i == 0 else '',
                item['item'],
                item['spec'],
                item['category'],
                float(item['unit_price'].replace('¥', '')) if item['unit_price'] else 0,
                int(item['quantity']),
                float(item['subtotal'].replace('¥', '')) if item['subtotal'] else 0,
                group['total'] if i == 0 else ''
            ])
            for cell in ws2[row_num]:
                cell.border = thin_border
            row_num += 1

        if len(group['items']) > 1:
            merge_info.append((start_row, row_num - 1, 'A'))
            merge_info.append((start_row, row_num - 1, 'H'))

    for start, end, col in merge_info:
        if start < end:
            ws2.merge_cells(f'{col}{start}:{col}{end}')
            cell = ws2[f'{col}{start}']
            cell.alignment = Alignment(horizontal='center', vertical='center')

    col_widths = [16, 24, 18, 14, 12, 8, 12, 16]
    for i, width in enumerate(col_widths):
        ws2.column_dimensions[get_column_letter(i + 1)].width = width

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    filename = f'quotation_{quotation_id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.xlsx'
    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


# ==================== 共享 PDF 渲染函数 ====================

def _build_pdf_tables(quotation, modules, fees, labor_hours,
                      packing_entries, person_days_entries, person_trip_entries, coeff):
    """构建PDF三个合并表格的数据"""
    rate_large = float(coeff.get('large', 1.0))
    rate_standard = float(coeff.get('standard', 1.0))
    rate_other = float(coeff.get('other', 1.0))
    is_line = (quotation.type == 'line')

    # ===== 表格1：设备硬件合计 =====
    table1_rows = []
    table1_total = 0.0

    if is_line:
        from collections import defaultdict
        modules_by_quotation = defaultdict(list)
        for mod in modules:
            modules_by_quotation[mod.quotation_id].append(mod)

        child_names = {}
        for c in quotation.children.all():
            child_names[c.id] = c.name

        for qid, mod_list in modules_by_quotation.items():
            group_subtotal = 0.0
            for mod in mod_list:
                if not hasattr(mod, 'materials') or not mod.materials:
                    continue
                for mm in mod.materials:
                    if not mm.material:
                        continue
                    rate = (rate_large if mm.material.category == 'large'
                            else (rate_standard if mm.material.category == 'standard' else rate_other))
                    if mm.is_other:
                        unit_amount = float(mm.unit_price_override or float(mm.material.unit_price or 0)) * rate
                    else:
                        unit_amount = float(mm.material.unit_price or 0) * rate * mm.quantity
                    group_subtotal += unit_amount

            display_name = '(线体)' if qid == quotation.id else child_names.get(qid, f'子项目{qid}')
            table1_rows.append({
                'module_name': display_name,
                'brand': 'RS',
                'quantity': 1,
                'unit': 'SET',
                'subtotal': round(group_subtotal, 2),
            })
            table1_total += group_subtotal
    else:
        for module in modules:
            if not module.materials:
                continue
            non_other_materials = [mm for mm in module.materials if mm.material and not mm.is_other]
            other_material = next((mm for mm in module.materials if mm.is_other), None)

            if len(non_other_materials) == 0:
                brand = '其他'
            elif len(non_other_materials) == 1:
                brand = non_other_materials[0].material.brand or ''
            else:
                brand = 'RS'

            module_subtotal = 0.0
            for mm in module.materials:
                if not mm.material:
                    continue
                rate = (rate_large if mm.material.category == 'large'
                        else (rate_standard if mm.material.category == 'standard' else rate_other))
                if mm.is_other:
                    unit_amount = float(mm.unit_price_override or float(mm.material.unit_price or 0)) * rate
                else:
                    unit_amount = float(mm.material.unit_price or 0) * rate * mm.quantity
                module_subtotal += unit_amount

            table1_rows.append({
                'module_name': module.name or '',
                'brand': brand,
                'quantity': 1,
                'unit': 'SET',
                'subtotal': round(module_subtotal, 2),
            })
            table1_total += module_subtotal

    # ===== 表格2：设备人力合计 =====
    table2_rows = []
    table2_total = 0.0
    for l in labor_hours:
        row_total = float(l.total or 0)
        table2_rows.append({
            'division': l.name or '',
            'hours': float(l.hours) if l.hours else 0,
            'unit': 'H',
            'hourly_rate': float(l.unit_price) if l.unit_price else 0,
            'subtotal': round(row_total, 2),
        })
        table2_total += row_total

    # ===== 三大新费用汇总 =====
    total_person_days = 0.0
    total_person_days_count = 0.0
    for entry in person_days_entries:
        up = float(entry.unit_price) if entry.unit_price else (
            float(entry.travel_category.day_rates[0].unit_price)
            if entry.travel_category and entry.travel_category.day_rates else 0)
        days = float(entry.person_days or 0)
        total_person_days += up * days
        total_person_days_count += days

    total_person_trips = 0.0
    total_person_trips_count = 0.0
    for entry in person_trip_entries:
        if entry.unit_price or entry.visa_fee:
            up = float(entry.unit_price) if entry.unit_price else 0
            vf = float(entry.visa_fee) if entry.visa_fee else 0
        else:
            fee_record = TravelPersonTripFee.query.filter_by(
                travel_category_id=entry.travel_category_id,
                travel_mode_id=entry.travel_mode_id,
                is_active=True
            ).first()
            up = float(fee_record.unit_price or 0) if fee_record else 0
            vf = float(fee_record.visa_fee or 0) if fee_record else 0
        cat_code = entry.travel_category.code if entry.travel_category else ''
        count = float(entry.person_count or 0)
        total_person_trips += count * (up + (vf if cat_code != 'domestic' else 0))
        total_person_trips_count += count

    total_packing = 0.0
    total_packing_count = 0.0
    for entry in packing_entries:
        up = float(entry.unit_price) if entry.unit_price else (
            float(entry.packing_type.unit_price)
            if entry.packing_type and entry.packing_type.unit_price else 0)
        qty = float(entry.quantity or 0)
        total_packing += up * qty
        total_packing_count += qty

    # ===== 表格3：其他合计 =====
    table3_rows = []
    table3_total = 0.0

    if total_person_days > 0:
        table3_rows.append({
            'name': '差旅住宿费',
            'quantity': round(total_person_days_count, 1),
            'unit': '人天',
            'unit_price': round(total_person_days, 2),
        })
        table3_total += total_person_days

    if total_person_trips > 0:
        table3_rows.append({
            'name': '差旅交通签证费',
            'quantity': round(total_person_trips_count, 0),
            'unit': '人次',
            'unit_price': round(total_person_trips, 2),
        })
        table3_total += total_person_trips

    if total_packing > 0:
        table3_rows.append({
            'name': '设备包装运输费',
            'quantity': round(total_packing_count, 0),
            'unit': '单元',
            'unit_price': round(total_packing, 2),
        })
        table3_total += total_packing

    dynamic_fees_total = 0.0
    for fee in fees:
        fee_amount = float(fee.amount or 0)
        table3_rows.append({
            'name': fee.fee_type or '',
            'quantity': 1,
            'unit': 'SET',
            'unit_price': round(fee_amount, 2),
        })
        dynamic_fees_total += fee_amount
    table3_total += dynamic_fees_total

    base_table3_total = table3_total

    profit_rate = float(quotation.profit_rate) if quotation.profit_rate else 0.0
    profit_amount = (table1_total + table2_total + table3_total) * profit_rate

    subtotal_with_profit = table1_total + table2_total + table3_total + profit_amount
    tax_rate = float(quotation.tax_rate) if quotation.tax_rate else 0.0
    tax_amount = subtotal_with_profit * tax_rate

    table3_rows.append({
        'name': '项目利润',
        'quantity': 1,
        'unit': 'SET',
        'unit_price': round(profit_amount, 2),
    })
    table3_rows.append({
        'name': '项目税额',
        'quantity': 1,
        'unit': 'SET',
        'unit_price': round(tax_amount, 2),
    })
    table3_total += profit_amount
    table3_total += tax_amount

    grand_total = subtotal_with_profit + tax_amount
    fees_subtotal = table3_total

    return {
        'table1': {'rows': table1_rows, 'total': round(table1_total, 2)},
        'table2': {'rows': table2_rows, 'total': round(table2_total, 2)},
        'table3': {'rows': table3_rows, 'total': round(table3_total, 2)},
        'grand_total': round(grand_total, 2),
        'fees_subtotal': round(fees_subtotal, 2),
        'profit_amount': round(profit_amount, 2),
        'tax_amount': round(tax_amount, 2),
        'subtotal_with_profit': round(subtotal_with_profit, 2),
        'fees_total': round(table3_total, 2),
        'rates': coeff,
    }


def render_pdf_bytes(tables, currency, currency_symbol, material_with_rates, table2_total, table3_total,
                     profit_rate, subtotal_with_profit, tax_rate, tax_amount, grand_total_converted,
                     quotation, data, lang):
    """共享 PDF 渲染逻辑，中英文通用"""
    class PDF(FPDF):
        def header(self): pass
        def footer(self): pass

    pdf = PDF()
    pdf.add_font('SimHei', '', FONT_REGULAR, uni=True)
    pdf.add_font('SimHei', 'B', FONT_BOLD, uni=True)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # 标题
    pdf.set_font('SimHei', 'B', 16)
    pdf.cell(0, 12, t('quote', lang), 0, 1, 'C')
    pdf.ln(5)

    # 基本信息
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('basic_info', lang), 0, 1)
    pdf.set_font('SimHei', '', 10)

    q_name = quotation.name if quotation else data.get('name', '')
    q_scheme_no = quotation.scheme_no if quotation else data.get('scheme_no', '')
    q_owner = quotation.business_owner.real_name if quotation and quotation.business_owner else ''
    q_created = quotation.created_at.strftime('%Y-%m-%d %H:%M') if quotation and quotation.created_at else ''
    q_version = data.get('version_no', '') if data else ''

    basic_info = [
        (t('project_name', lang), q_name or ''),
        (t('scheme_no', lang), q_scheme_no or ''),
        (t('business_owner', lang), q_owner),
        (t('created_at', lang), q_created),
        (t('currency', lang), currency),
    ]
    if q_version:
        basic_info.append((t('version_no', lang), f'v{q_version}'))

    for key, value in basic_info:
        pdf.set_font('SimHei', 'B', 9)
        pdf.cell(40, 6, key, 1)
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 6, str(value), 1, 1)

    pdf.ln(5)

    # 合并表格
    if tables['table1']['rows'] or tables['table2']['rows'] or tables['table3']['rows']:
        unified_col_widths = [60, 35, 30, 20, 45]

        # 1. 设备硬件
        if tables['table1']['rows']:
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(sum(unified_col_widths), 5, '【设备硬件】', 1, 1, 'L', True)
            pdf.set_fill_color(217, 217, 217)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(unified_col_widths[0], 6, t('module_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[1], 6, t('brand_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[2], 6, t('qty_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[3], 6, t('unit_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, t('total_col', lang), 1, 1, 'C', True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
            for row in tables['table1']['rows']:
                pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
                pdf.cell(unified_col_widths[0], 5, str(row['module_name'])[:20], 1)
                pdf.cell(unified_col_widths[1], 5, str(row['brand'])[:10], 1)
                qty_val = row['quantity']
                pdf.cell(unified_col_widths[2], 5, str(int(qty_val)) if qty_val == int(qty_val) else str(qty_val), 1, 0, 'C')
                pdf.cell(unified_col_widths[3], 5, str(row['unit']), 1, 0, 'C')
                pdf.cell(unified_col_widths[4], 5, f'¥{row["subtotal"]:.2f}', 1, 1, 'R')
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(226, 239, 218)
            pdf.cell(sum(unified_col_widths[:4]), 6, t('subtotal_label', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, f'¥{tables["table1"]["total"]:.2f}', 1, 1, 'R', True)

        # 2. 设备人力
        if tables['table2']['rows']:
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(sum(unified_col_widths), 5, '【设备人力】', 1, 1, 'L', True)
            pdf.set_fill_color(217, 217, 217)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(unified_col_widths[0] + unified_col_widths[1], 6, t('division_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[2], 6, t('qty_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[3], 6, t('unit_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, t('total_col', lang), 1, 1, 'C', True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
            for row in tables['table2']['rows']:
                pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
                pdf.cell(unified_col_widths[0] + unified_col_widths[1], 5, str(row['division'])[:20], 1)
                pdf.cell(unified_col_widths[2], 5, f'{row["hours"]:.1f}', 1, 0, 'C')
                pdf.cell(unified_col_widths[3], 5, 'H', 1, 0, 'C')
                pdf.cell(unified_col_widths[4], 5, f'¥{row["subtotal"]:.2f}', 1, 1, 'R')
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(226, 239, 218)
            pdf.cell(sum(unified_col_widths[:4]), 6, t('subtotal_label', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, f'¥{tables["table2"]["total"]:.2f}', 1, 1, 'R', True)

        # 3. 其他合计
        if tables['table3']['rows']:
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(sum(unified_col_widths), 5, '【其他】', 1, 1, 'L', True)
            pdf.set_fill_color(217, 217, 217)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(unified_col_widths[0] + unified_col_widths[1], 6, t('fee_name_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[2], 6, t('qty_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[3], 6, t('unit_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, t('unit_price_col', lang), 1, 1, 'C', True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
            for row in tables['table3']['rows']:
                pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
                pdf.cell(unified_col_widths[0] + unified_col_widths[1], 5, str(row['name'])[:20], 1)
                qty_val = row.get('quantity', 1)
                pdf.cell(unified_col_widths[2], 5, str(int(qty_val)) if isinstance(qty_val, (int, float)) and qty_val == int(qty_val) else str(qty_val), 1, 0, 'C')
                pdf.cell(unified_col_widths[3], 5, str(row.get('unit', 'SET')), 1, 0, 'C')
                pdf.cell(unified_col_widths[4], 5, f'¥{row.get("unit_price", 0):.2f}', 1, 1, 'R')
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(226, 239, 218)
            pdf.cell(sum(unified_col_widths[:4]), 6, t('subtotal_label', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, f'¥{tables["fees_subtotal"]:.2f}', 1, 1, 'R', True)

        # 最终报价
        pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
        pdf.set_font('SimHei', 'B', 9)
        pdf.set_fill_color(255, 255, 204)
        pdf.cell(sum(unified_col_widths[:2]), 6, t('grand_total', lang), 1, 0, 'C', True)
        pdf.cell(sum(unified_col_widths[2:4]), 6, currency, 1, 0, 'C', True)
        pdf.cell(unified_col_widths[4], 6, f'{currency_symbol}{grand_total_converted:.2f}', 1, 1, 'R', True)

    return pdf.output()


# ==================== 导出 PDF（实时） ====================

@router.get('/quotations/{quotation_id}/export/pdf')
def export_pdf(
    quotation_id: int,
    lang: str = Query('zh', description='语言: zh/en'),
    db=Depends(get_db),
):
    """导出报价单 PDF 格式，支持中英文 ?lang=en"""
    quotation = db.query(Quotation).get(quotation_id)
    if not quotation:
        raise HTTPException(status_code=404, detail='Quotation not found')

    # 线体：聚合所有子报价单数据；普通：只查自己（与 Flask 一致）
    if quotation.type == 'line':
        child_ids = [c.id for c in quotation.children.all()]
        all_ids = [quotation_id] + child_ids
        modules = Module.query.filter(Module.quotation_id.in_(all_ids)).all()
        for module in modules:
            module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()
        fees = OtherFee.query.filter(OtherFee.quotation_id.in_(all_ids)).all()
        labor_hours = LaborHour.query.filter(LaborHour.quotation_id.in_(all_ids)).all()
        packing_entries = PackingEntry.query.filter(PackingEntry.quotation_id.in_(all_ids)).all()
        person_days_entries = TravelPersonDays.query.filter(TravelPersonDays.quotation_id.in_(all_ids)).all()
        person_trip_entries = TravelPersonTrip.query.filter(TravelPersonTrip.quotation_id.in_(all_ids)).all()
    else:
        modules = Module.query.filter_by(quotation_id=quotation_id).all()
        for module in modules:
            module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()
        fees = OtherFee.query.filter_by(quotation_id=quotation_id).all()
        labor_hours = LaborHour.query.filter_by(quotation_id=quotation_id).all()
        packing_entries = PackingEntry.query.filter_by(quotation_id=quotation_id).all()
        person_days_entries = TravelPersonDays.query.filter_by(quotation_id=quotation_id).all()
        person_trip_entries = TravelPersonTrip.query.filter_by(quotation_id=quotation_id).all()

    coeff = (quotation.coefficients if hasattr(quotation, 'coefficients') and quotation.coefficients else
             {fr.category: float(fr.rate) for fr in FeeRate.query.all()})

    tables = _build_pdf_tables(quotation, modules, fees, labor_hours,
                               packing_entries, person_days_entries, person_trip_entries, coeff)

    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    currency_symbol = get_currency_symbol(currency)

    profit_rate = float(quotation.profit_rate) if quotation.profit_rate else 0.0
    tax_rate = float(quotation.tax_rate) if quotation.tax_rate else 0.0
    grand_total = tables.get('subtotal_with_profit', 0) + tables.get('tax_amount', 0)
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    grand_total_converted = convert_currency(grand_total, currency, exchange_rates)

    pdf_bytes = render_pdf_bytes(
        tables, currency, currency_symbol,
        tables.get('table1', {}).get('total', 0), tables.get('table2', {}).get('total', 0), tables['fees_subtotal'],
        profit_rate, tables.get('subtotal_with_profit', 0), tax_rate,
        tables.get('tax_amount', 0), grand_total_converted,
        quotation, None, lang
    )

    buffer = BytesIO(pdf_bytes)
    filename = f'quotation_{quotation_id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.pdf'
    return StreamingResponse(
        buffer,
        media_type='application/pdf',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


# ==================== 版本导出（原地生成，不再代理 Flask 5000） ====================

import os
from fastapi.responses import FileResponse

@router.get('/quotations/{quotation_id}/versions/{version_no}/export/{fmt}')
def export_version(
    quotation_id: int,
    version_no: int,
    fmt: str,
    lang: str = Query('zh'),
    db=Depends(get_db),
):
    """导出版本文件：从数据库 pdf_file/word_file 字段读取真实路径"""
    import json as _json
    from app.models.version import VersionSnapshot

    ver = db.query(VersionSnapshot).filter_by(
        quotation_id=quotation_id, version_no=version_no
    ).first()
    if not ver:
        raise HTTPException(status_code=404, detail=f'版本 v{version_no} 不存在')

    mime_map = {
        'pdf': 'application/pdf',
        'word': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'excel': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    }

    if fmt == 'pdf':
        # pdf_file 字段是 JSON 字符串 {"zh": "path", "en": "path"}
        pdf_paths = _json.loads(ver.pdf_file) if isinstance(ver.pdf_file, str) and ver.pdf_file else (ver.pdf_file or {})
        lang_key = 'zh' if lang != 'en' else 'en'
        filepath = pdf_paths.get(lang_key) or pdf_paths.get('zh')
        if not filepath:
            raise HTTPException(status_code=404, detail='版本 PDF 不存在，请先归档')
    elif fmt == 'word':
        filepath = ver.word_file
        if not filepath:
            raise HTTPException(status_code=404, detail='版本 Word 不存在')
    elif fmt == 'excel':
        raise HTTPException(status_code=400, detail='版本暂不支持 Excel 导出')
    else:
        raise HTTPException(status_code=400, detail=f'不支持的格式: {fmt}')

    if not os.path.exists(filepath):
        raise HTTPException(
            status_code=404,
            detail=f'文件已被清理或归档路径失效，请重新归档生成（路径: {filepath}）'
        )

    return FileResponse(
        filepath,
        media_type=mime_map.get(fmt, 'application/octet-stream'),
        filename=os.path.basename(filepath),
    )
