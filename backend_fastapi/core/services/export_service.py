"""导出功能模块 - 支持 Word/Excel/PDF 格式（业务函数，无 Flask 装饰器）"""
from io import BytesIO
from datetime import datetime

# Word 导出
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Excel 导出
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# PDF 导出 - 使用 fpdf2 支持中文
from fpdf import FPDF

# 中文字体路径（相对路径，字体跟随项目）
import os
# fonts 在 backend_fastapi/fonts/，export_service.py 在 core/services/
# 即 __file__ 向上 2 级：services → core → backend_fastapi
FONT_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'fonts')
FONT_REGULAR = os.path.join(FONT_DIR, 'simhei.ttf')  # 黑体（支持中文）
FONT_BOLD = os.path.join(FONT_DIR, 'simhei.ttf')       # 黑体（暂无粗体，用黑体代替）

# 中英文标签映射
I18N = {
    'zh': {
        'quote': '报价单', 'basic_info': '基本信息', 'project_name': '项目名称',
        'scheme_no': '方案编号', 'business_owner': '业务负责人', 'created_at': '创建时间',
        'currency': '报价币种', 'material_list': '物料及费用清单',
        'module_col': '模块', 'item_col': '项目', 'spec_col': '规格',
        'category_col': '分类/品牌', 'unit_price_col': '单价', 'qty_col': '数量',
        'subtotal_col': '小计', 'total_col': '合计',
        'quote_summary': '报价汇总', 'project_col': '项目', 'amount_col': '金额',
        'material_total': '物料合计', 'fee_total': '费用合计', 'labor_total': '人力合计',
        'profit_rate': '利润率', 'subtotal_with_profit': '含利润率小计',
        'tax_rate': '税率', 'tax_amount': '税额', 'grand_total': '最终报价',
        'other_fees': '其他费用', 'unnamed_module': '未命名模块',
        'quote_summary': '报价汇总',
        'internal': '厂内', 'external': '厂外',
        'no_materials': '（无物料）', 'no_data': '无物料及费用数据',
        'labor_hours': '人力工时', 'version': '版本', 'version_no': '版本号',
        'profit_rate_label': '利润率', 'subtotal_label': '合计',
        'brand_col': '品牌', 'division_col': '设计工时划分', 'fee_name_col': '费用名称',
        'table1_total': '设备硬件合计', 'table2_total': '设备人力合计', 'table3_total': '其他合计',
        'profit_amount': '项目利润',
        'unit_col': '单位', 'total_col': '合计',
        'unit_h': 'h', 'tax_included': '含税报价',
    },
    'en': {
        'quote': 'QUOTATION', 'basic_info': 'Basic Information', 'project_name': 'Project Name',
        'scheme_no': 'Scheme No.', 'business_owner': 'Sales Owner', 'created_at': 'Created At',
        'currency': 'Currency', 'material_list': 'Material & Cost List',
        'module_col': 'Module', 'item_col': 'Item', 'spec_col': 'Spec',
        'category_col': 'Category/Brand', 'unit_price_col': 'Unit Price', 'qty_col': 'Qty',
        'subtotal_col': 'Subtotal', 'total_col': 'Total',
        'quote_summary': 'Quote Summary', 'project_col': 'Item', 'amount_col': 'Amount',
        'material_total': 'Material Total', 'fee_total': 'Fee Total', 'labor_total': 'Labor Total',
        'profit_rate': 'Profit Rate', 'subtotal_with_profit': 'Subtotal w/ Profit',
        'tax_rate': 'Tax Rate', 'tax_amount': 'Tax Amount', 'grand_total': 'Grand Total',
        'other_fees': 'Other Fees', 'unnamed_module': 'Unnamed Module',
        'quote_summary': 'Quote Summary',
        'internal': 'Internal', 'external': 'External',
        'no_materials': '(No Materials)', 'no_data': 'No material or cost data',
        'labor_hours': 'Labor Hours', 'version': 'Version', 'version_no': 'Version No.',
        'profit_rate_label': 'Profit Rate', 'subtotal_label': 'Subtotal w/ Profit',
        'table1_total': 'Hardware Total', 'table2_total': 'Labor Total', 'table3_total': 'Other Total',
        'profit_amount': 'Project Profit',
        'unit_col': 'Unit', 'total_col': 'Total',
        'unit_h': 'h', 'tax_included': 'Tax Included',
    }
}

from core.models.quotation import Quotation
from core.models.module import Module
from core.models.material import ModuleMaterial, Material
from core.models.fee import OtherFee, FeeType
from core.models.fee_rate import FeeRate
from core.models.exchange_rate import ExchangeRate
from core.models.version import VersionSnapshot
from core.models.user import User
from core.models.labor_hour import LaborHour
from core.models.travel_entry import PackingEntry, TravelPersonDays, TravelPersonTrip
from core.models.travel import TravelPersonTripFee
from core.models.packing import PackingType
from db import db


def t(key, lang='zh'):
    return I18N.get(lang, I18N['zh']).get(key, key)


def get_name(obj, lang='zh'):
    """获取对象的本地化名称，英文时优先用 name_en"""
    if lang == 'en':
        en = getattr(obj, 'name_en', None)
        if en:
            return en
    return obj.name


def get_fee_type_name(fee_type_id, lang='zh'):
    """根据 fee_type ID 或名称字符串获取费用类型名称，英文时优先用 name_en"""
    if not fee_type_id:
        return fee_type_id or ''
    # 如果是整数ID，按ID查
    if isinstance(fee_type_id, int):
        ft = db.session.get(FeeType, fee_type_id)
        if ft:
            return ft.name_en if (lang == 'en' and ft.name_en) else ft.name
        return str(fee_type_id)
    # 如果是字符串（可能是中文名或英文名），按名称查
    ft = FeeType.query.filter(FeeType.name == str(fee_type_id)).first()
    if ft:
        return ft.name_en if (lang == 'en' and ft.name_en) else ft.name
    # 也尝试查 name_en
    if lang == 'en':
        ft = FeeType.query.filter(FeeType.name_en == str(fee_type_id)).first()
        if ft:
            return ft.name_en
    return str(fee_type_id)


# ==================== 数据获取 ====================


def get_quotation_with_details(quotation_id):
    """获取报价单及其详细信息（含三大新费用）
    线体报价单：聚合所有子报价单的数据
    """
    quotation = Quotation.query.get(quotation_id)
    if not quotation:
        return None

    # 线体：聚合所有子报价单的数据
    if quotation.type == 'line':
        child_ids = [c.id for c in quotation.children.all()]
        all_ids = [quotation_id] + child_ids

        # 所有模块（含子报价单的模块）
        modules = Module.query.filter(Module.quotation_id.in_(all_ids)).all()
        for module in modules:
            module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()

        # 费用聚合
        fees = OtherFee.query.filter(OtherFee.quotation_id.in_(all_ids)).all()
        labor_hours = LaborHour.query.filter(LaborHour.quotation_id.in_(all_ids)).all()
        packing_entries = PackingEntry.query.filter(PackingEntry.quotation_id.in_(all_ids)).all()
        person_days_entries = TravelPersonDays.query.filter(TravelPersonDays.quotation_id.in_(all_ids)).all()
        person_trip_entries = TravelPersonTrip.query.filter(TravelPersonTrip.quotation_id.in_(all_ids)).all()
    else:
        # 普通报价单
        modules = Module.query.filter_by(quotation_id=quotation_id).all()
        for module in modules:
            module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()
        fees = OtherFee.query.filter_by(quotation_id=quotation_id).all()
        labor_hours = LaborHour.query.filter_by(quotation_id=quotation_id).all()
        packing_entries = PackingEntry.query.filter_by(quotation_id=quotation_id).all()
        person_days_entries = TravelPersonDays.query.filter_by(quotation_id=quotation_id).all()
        person_trip_entries = TravelPersonTrip.query.filter_by(quotation_id=quotation_id).all()

    return quotation, modules, fees, labor_hours, packing_entries, person_days_entries, person_trip_entries


# ==================== 表格构建 ====================


def _build_pdf_tables(quotation, modules, fees, labor_hours, packing_entries, person_days_entries, person_trip_entries, coeff):
    """构建PDF三个合并表格的数据
    line报价单：设备硬件按子项目分组汇总
    """
    rate_large = float(coeff.get('large', 1.0))
    rate_standard = float(coeff.get('standard', 1.0))
    rate_other = float(coeff.get('other', 1.0))
    is_line = (quotation.type == 'line')

    # ===== 表格1：设备硬件合计 =====
    table1_rows = []
    table1_total = 0.0

    if is_line:
        # 线体模式：按子报价单分组，每个子报价单的物料×系数后合计
        # 先按 quotation_id 分组
        from collections import defaultdict
        modules_by_quotation = defaultdict(list)
        for mod in modules:
            modules_by_quotation[mod.quotation_id].append(mod)

        # 子报价单名称映射
        child_names = {}
        for c in quotation.children.all():
            child_names[c.id] = c.name

        for qid, mod_list in modules_by_quotation.items():
            # 该组的含系数合计
            group_subtotal = 0.0
            for mod in mod_list:
                if not hasattr(mod, 'materials') or not mod.materials:
                    continue
                for mm in mod.materials:
                    # migration 020: 自制件用 mm.category 选 rate
                    if mm.is_custom:
                        cat = mm.category or 'standard'
                        rate = (rate_large if cat == 'large'
                                else (rate_standard if cat == 'standard' else rate_other))
                        unit_amount = float(mm.unit_price_override or 0) * rate * float(mm.quantity or 1)
                        group_subtotal += unit_amount
                        continue
                    if not mm.material:
                        continue
                    rate = (rate_large if mm.material.category == 'large'
                            else (rate_standard if mm.material.category == 'standard' else rate_other))
                    if mm.is_other:
                        unit_amount = float(mm.unit_price_override or float(mm.material.unit_price or 0)) * rate
                    else:
                        unit_amount = float(mm.material.unit_price or 0) * rate * mm.quantity
                    group_subtotal += unit_amount

            # 确定显示名称
            if qid == quotation.id:
                display_name = '(线体)'
            else:
                display_name = child_names.get(qid, f'子项目{qid}')

            table1_rows.append({
                'module_name': display_name,
                'brand': 'RS',
                'quantity': 1,
                'unit': 'SET',
                'subtotal': round(group_subtotal, 2),
            })
            table1_total += group_subtotal
    else:
        # 普通模式：按模块分列
        for module in modules:
            if not module.materials:
                continue
            non_other_materials = [mm for mm in module.materials
                                  if mm.material and not mm.is_other]
            other_material = next((mm for mm in module.materials
                                   if mm.is_other), None)

            # 品牌取值
            if len(non_other_materials) == 0:
                # 只有其他或无物料
                brand = '其他'
            elif len(non_other_materials) == 1:
                # 只有一个非其他物料
                brand = non_other_materials[0].material.brand or ''
            else:
                # 多个非其他物料
                brand = 'RS'

            # 该模块的含系数小计
            module_subtotal = 0.0
            for mm in module.materials:
                # migration 020: 自制件用 mm.category 选 rate
                if mm.is_custom:
                    cat = mm.category or 'standard'
                    rate = (rate_large if cat == 'large'
                            else (rate_standard if cat == 'standard' else rate_other))
                    unit_amount = float(mm.unit_price_override or 0) * rate * float(mm.quantity or 1)
                    module_subtotal += unit_amount
                    continue
                if not mm.material:
                    continue
                rate = (rate_large if mm.material.category == 'large'
                        else (rate_standard if mm.material.category == 'standard' else rate_other))
                if mm.is_other:
                    unit_amount = float(mm.unit_price_override or float(mm.material.unit_price or 0)) * rate
                else:
                    unit_amount = float(mm.material.unit_price or 0) * rate * mm.quantity
                module_subtotal += unit_amount

            table1_rows.append({
                'module_name': module.name or '',
                'brand': brand,
                'quantity': 1,
                'unit': 'SET',
                'subtotal': round(module_subtotal, 2),
            })
            table1_total += module_subtotal

    # ===== 表格2：设备人力合计 =====
    table2_rows = []
    table2_total = 0.0
    for l in labor_hours:
        row_total = float(l.total or 0)
        table2_rows.append({
            'division': l.name or '',
            'hours': float(l.hours) if l.hours else 0,
            'unit': 'H',
            'hourly_rate': float(l.unit_price) if l.unit_price else 0,
            'subtotal': round(row_total, 2),
        })
        table2_total += row_total

    # ===== 三大新费用汇总 =====
    # 差旅人天费用 + 总人天数量
    total_person_days = 0.0
    total_person_days_count = 0.0
    for entry in person_days_entries:
        up = float(entry.unit_price) if entry.unit_price else (
            float(entry.travel_category.day_rates[0].unit_price)
            if entry.travel_category and entry.travel_category.day_rates else 0)
        days = float(entry.person_days or 0)
        total_person_days += up * days
        total_person_days_count += days

    # 差旅人次费用 + 总人次数量
    total_person_trips = 0.0
    total_person_trips_count = 0.0
    for entry in person_trip_entries:
        if entry.unit_price or entry.visa_fee:
            up = float(entry.unit_price) if entry.unit_price else 0
            vf = float(entry.visa_fee) if entry.visa_fee else 0
        else:
            fee_record = TravelPersonTripFee.query.filter_by(
                travel_category_id=entry.travel_category_id,
                travel_mode_id=entry.travel_mode_id,
                is_active=True
            ).first()
            up = float(fee_record.unit_price or 0) if fee_record else 0
            vf = float(fee_record.visa_fee or 0) if fee_record else 0
        cat_code = entry.travel_category.code if entry.travel_category else ''
        count = float(entry.person_count or 0)
        total_person_trips += count * (up + (vf if cat_code != 'domestic' else 0))
        total_person_trips_count += count

    # 运输包装费用 + 总数量
    total_packing = 0.0
    total_packing_count = 0.0
    for entry in packing_entries:
        up = float(entry.unit_price) if entry.unit_price else (
            float(entry.packing_type.unit_price)
            if entry.packing_type and entry.packing_type.unit_price else 0)
        qty = float(entry.quantity or 0)
        total_packing += up * qty
        total_packing_count += qty

    # ===== 表格3：其他合计 =====
    table3_rows = []
    table3_total = 0.0

    # 固定条目：差旅住宿费
    if total_person_days > 0:
        table3_rows.append({
            'name': '差旅住宿费',
            'quantity': round(total_person_days_count, 1),
            'unit': '人天',
            'unit_price': round(total_person_days, 2),
        })
        table3_total += total_person_days

    # 固定条目：差旅交通签证费
    if total_person_trips > 0:
        table3_rows.append({
            'name': '差旅交通签证费',
            'quantity': round(total_person_trips_count, 0),
            'unit': '人次',
            'unit_price': round(total_person_trips, 2),
        })
        table3_total += total_person_trips

    # 固定条目：设备包装运输费
    if total_packing > 0:
        table3_rows.append({
            'name': '设备包装运输费',
            'quantity': round(total_packing_count, 0),
            'unit': '单元',
            'unit_price': round(total_packing, 2),
        })
        table3_total += total_packing

    # 动态条目：费用Tab每项
    dynamic_fees_total = 0.0
    for fee in fees:
        fee_amount = float(fee.amount or 0)
        table3_rows.append({
            'name': fee.fee_type or '',
            'quantity': 1,
            'unit': 'SET',
            'unit_price': round(fee_amount, 2),
        })
        dynamic_fees_total += fee_amount
    table3_total += dynamic_fees_total

    # 基础 table3_total（不含利润和税额行）
    base_table3_total = table3_total

    # 项目利润 = (物料×系数 + 人力合计 + 基础费用) × 对外利润率
    profit_rate = float(quotation.profit_rate) if quotation.profit_rate else 0.0
    profit_amount = (table1_total + table2_total + table3_total) * profit_rate

    # 项目税额
    subtotal_with_profit = table1_total + table2_total + table3_total + profit_amount
    tax_rate = float(quotation.tax_rate) if quotation.tax_rate else 0.0
    tax_amount = subtotal_with_profit * tax_rate

    table3_rows.append({
        'name': '项目利润',
        'quantity': 1,
        'unit': 'SET',
        'unit_price': round(profit_amount, 2),
    })
    table3_rows.append({
        'name': '项目税额',
        'quantity': 1,
        'unit': 'SET',
        'unit_price': round(tax_amount, 2),
    })
    table3_total += profit_amount
    table3_total += tax_amount

    grand_total = subtotal_with_profit + tax_amount

    # fees_subtotal = 其他合计 + 项目利润 + 税额
    fees_subtotal = table3_total

    return {
        'table1': {'rows': table1_rows, 'total': round(table1_total, 2)},
        'table2': {'rows': table2_rows, 'total': round(table2_total, 2)},
        'table3': {'rows': table3_rows, 'total': round(table3_total, 2)},
        'grand_total': round(grand_total, 2),
        'fees_subtotal': round(fees_subtotal, 2),
        'profit_amount': round(profit_amount, 2),
        'tax_amount': round(tax_amount, 2),
        'subtotal_with_profit': round(subtotal_with_profit, 2),
        'fees_total': round(table3_total, 2),
        'rates': coeff,
    }


# ==================== 汇率/币种工具函数 ====================


def convert_currency(amount, to_currency, exchange_rates):
    """汇率转换"""
    if to_currency == 'CNY':
        return amount
    # rate 是 "1 目标货币 = X CNY"
    rate = exchange_rates.get(to_currency, 1)
    return amount / rate


def get_currency_symbol(currency):
    """获取币种符号"""
    symbols = {'CNY': '¥', 'USD': '$', 'EUR': '€', 'GBP': '£', 'JPY': '¥'}
    return symbols.get(currency, '¥')


# ==================== 版本文件生成 ====================


def generate_version_files(quotation_id, version_no, snapshot_data, data=None, lang='zh'):
    """为版本生成Word和PDF文件，返回文件路径字典"""
    import os
    import json

    # 创建版本文件存储目录（统一存放在 backend_fastapi/static/versions/）
    # __file__ = core/services/export_service.py → 3 次 dirname 升到 backend_fastapi
    base_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'static', 'versions', f'quotation_{quotation_id}')
    os.makedirs(base_dir, exist_ok=True)

    word_path = os.path.join(base_dir, f'v{version_no}.docx')
    pdf_zh_path = os.path.join(base_dir, f'v{version_no}_zh.pdf')
    pdf_en_path = os.path.join(base_dir, f'v{version_no}_en.pdf')

    # 获取报价单信息（包含币种）
    quotation = Quotation.query.get(quotation_id)

    if data is None:
        data = json.loads(snapshot_data) if isinstance(snapshot_data, str) else snapshot_data

    # 填充物料详细信息（与 get_version_snapshot_data 相同逻辑）
    # 注意：保留 is_other 物料的 unit_price_override，不被 live 数据覆盖
    for module in data.get('modules', []):
        for mm in module.get('materials', []):
            mat = Material.query.get(mm.get('material_id'))
            if mat:
                mm['name'] = mat.name or ''
                mm['brand'] = mat.brand or ''
                mm['spec'] = mat.spec or ''
                # 只有非 is_other 物料才用 live 价；is_other 保留快照中的 override 价
                if not mm.get('is_other', False):
                    mm['unit_price'] = float(mat.unit_price or 0)
                mm['category'] = mat.category or 'standard'

    # 计算汇总
    totals = calculate_version_totals(data)

    # 生成 Word 文件
    _generate_version_word(word_path, quotation_id, version_no, data, totals, quotation, 'zh')

    # 生成 PDF 文件（中英文各一份）
    _generate_version_pdf(pdf_zh_path, quotation_id, version_no, data, totals, quotation, 'zh')
    _generate_version_pdf(pdf_en_path, quotation_id, version_no, data, totals, quotation, 'en')

    import json
    pdf_file_json = json.dumps({'zh': pdf_zh_path, 'en': pdf_en_path})
    return {'word': word_path, 'pdf': pdf_file_json}


# ==================== Word 生成 ====================


def _generate_version_word(word_path, quotation_id, version_no, data, totals, quotation=None, lang='zh'):
    """生成版本 Word 文件"""
    doc = Document()
    title = doc.add_heading(f'{data.get("name", t("quote", lang))}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 币种和汇率
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)

    # 基本信息
    doc.add_heading(t('basic_info', lang), level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    business_owner_name = ''
    if data.get('business_owner_id'):
        owner = User.query.get(data.get('business_owner_id'))
        if owner:
            business_owner_name = owner.real_name
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    info_data = [
        (t('project_name', lang), data.get('name', '')),
        (t('scheme_no', lang), data.get('scheme_no', '')),
        (t('business_owner', lang), business_owner_name),
        (t('version_no', lang), f'v{version_no}'),
        (t('currency', lang), currency)
    ]
    for i, (key, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # 整合表格
    doc.add_heading(t('material_list', lang), level=1)

    module_groups = []
    rate_large = totals.get('fee_rates', {}).get('large', 1.0)
    rate_standard = totals.get('fee_rates', {}).get('standard', 1.0)
    rate_other = totals.get('fee_rates', {}).get('other', 1.0)
    for module in data.get('modules', []):
        items = []
        module_total = 0
        for mm in module.get('materials', []):
            cat = mm.get('category', 'standard')
            rate = rate_large if cat == 'large' else (rate_standard if cat == 'standard' else rate_other)
            unit_price_with_rate = float(mm.get('unit_price', 0) or 0) * rate
            item_total = unit_price_with_rate * float(mm.get('quantity', 0) or 0)
            module_total += item_total
            items.append({
                'item': mm.get('name', ''),
                'spec': mm.get('spec', ''),
                'category': mm.get('brand', ''),
                'unit_price': f'¥{unit_price_with_rate:.2f}',
                'quantity': str(mm.get('quantity', 0)),
                'subtotal': f'¥{item_total:.2f}',
            })
        module_groups.append({
            'name': module.get('name', t('unnamed_module', lang)),
            'items': items,
            'total': module_total
        })

    fees_total = sum(float(f.get('amount', 0)) for f in data.get('fees', []))
    fee_items = []
    for fee in data.get('fees', []):
        location_text = t('internal', lang) if fee.get('location') == 'internal' else (t('external', lang) if fee.get('location') == 'external' else '')
        fee_items.append({
            'item': get_fee_type_name(fee.get('fee_type', ''), lang) or fee.get('name', ''),
            'spec': location_text,
            'category': get_fee_type_name(fee.get('fee_type', ''), lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f'¥{float(fee.get("amount", 0)):.2f}',
        })
    # 人力工时明细加入"其他费用"分组
    labor_total_sum = 0.0
    for lh in data.get('labor_hours', []):
        labor_total_sum += float(lh.get('total', 0))
        fee_items.append({
            'item': lh.get('name', '人力工时'),
            'spec': '',
            'category': t('labor_hours', lang),
            'unit_price': f'¥{float(lh.get("unit_price", 0)):.2f}',
            'quantity': str(lh.get('hours', 0)),
            'subtotal': f'¥{float(lh.get("total", 0)):.2f}',
        })
    fees_total += labor_total_sum
    if fee_items:
        module_groups.append({
            'name': t('other_fees', lang),
            'items': fee_items,
            'total': fees_total
        })

    if module_groups:
        combined_table = doc.add_table(rows=1, cols=8)
        combined_table.style = 'Table Grid'
        headers = ['模块', '项目', '规格', '分类/品牌', '单价', '数量', '小计', '合计']
        for j, header in enumerate(headers):
            combined_table.rows[0].cells[j].text = header

        for group in module_groups:
            if not group['items']:
                row = combined_table.add_row()
                row.cells[0].text = group['name']
                row.cells[1].text = '（无物料）'
                continue

            first_row = combined_table.add_row()
            first_row.cells[0].text = group['name']
            first_row.cells[1].text = group['items'][0]['item']
            first_row.cells[2].text = group['items'][0]['spec']
            first_row.cells[3].text = group['items'][0]['category']
            first_row.cells[4].text = group['items'][0]['unit_price']
            first_row.cells[5].text = group['items'][0]['quantity']
            first_row.cells[6].text = group['items'][0]['subtotal']
            first_row.cells[7].text = f'¥{group["total"]:.2f}'

            if len(group['items']) > 1:
                module_cell = first_row.cells[0]
                for i in range(1, len(group['items'])):
                    next_row = combined_table.add_row()
                    next_row.cells[1].text = group['items'][i]['item']
                    next_row.cells[2].text = group['items'][i]['spec']
                    next_row.cells[3].text = group['items'][i]['category']
                    next_row.cells[4].text = group['items'][i]['unit_price']
                    next_row.cells[5].text = group['items'][i]['quantity']
                    next_row.cells[6].text = group['items'][i]['subtotal']
                    module_cell = module_cell.merge(next_row.cells[0])

                total_cell = first_row.cells[7]
                for i in range(1, len(group['items'])):
                    next_row_idx = len(combined_table.rows) - 1
                    total_cell = total_cell.merge(combined_table.rows[next_row_idx].cells[7])
    else:
        doc.add_paragraph('无物料及费用数据')

    doc.add_paragraph()

    # 汇总信息
    doc.add_heading(t('quote_summary', lang), level=1)
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Table Grid'
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary_data = [
        (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
        (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
        (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
        (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
        (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
        (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
        (f'{t("grand_total", lang)}({currency})', f'{currency_symbol}{grand_total_converted:.2f}')
    ]
    for i, (key, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = value

    doc.save(word_path)


# ==================== PDF 生成 ====================


def _generate_version_pdf(pdf_path, quotation_id, version_no, data, totals, quotation=None, lang='zh'):
    """生成版本 PDF 文件（复用报价单PDF逻辑）"""
    # 获取币种和汇率
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)

    # 获取费用系数（优先用快照中的报价单系数，否则用系统费率）
    fee_rates = data.get('coefficients') or totals.get('fee_rates', {})
    rate_large = fee_rates.get('large', 1.0)
    rate_standard = fee_rates.get('standard', 1.0)
    rate_other = fee_rates.get('other', 1.0)

    # 收集分组数据
    module_groups = []
    for module in data.get('modules', []):
        items = []
        module_total = 0
        for mm in module.get('materials', []):
            mat = Material.query.get(mm.get('material_id'))
            if mat:
                cat = mat.category or 'standard'
                rate = rate_large if cat == 'large' else (rate_standard if cat == 'standard' else rate_other)
                unit_price_with_rate = float(mat.unit_price or 0) * rate
                item_total = unit_price_with_rate * mm.get('quantity', 1)
                module_total += item_total
                items.append({
                    'item': mat.name or '',
                    'spec': mat.spec or '',
                    'category': mat.brand or '',
                    'unit_price': f'¥{unit_price_with_rate:.2f}',
                    'quantity': str(mm.get('quantity', 1)),
                    'subtotal': f'¥{item_total:.2f}',
                })
        if items:
            module_obj = Module.query.get(module.get('id'))
            module_name = get_name(module_obj, lang) if module_obj else module.get('name', t('unnamed_module', lang))
            module_groups.append({
                'name': module_name,
                'items': items,
                'total': module_total
            })

    fees_total = sum(float(f.get('amount', 0)) for f in data.get('fees', []))
    fee_items = []
    for fee in data.get('fees', []):
        location_text = '厂内' if fee.get('location') == 'internal' else ('厂外' if fee.get('location') == 'external' else '')
        fee_items.append({
            'item': get_fee_type_name(fee.get('fee_type', ''), lang) or fee.get('name', ''),
            'spec': location_text,
            'category': get_fee_type_name(fee.get('fee_type', ''), lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f"¥{float(fee.get('amount', 0)):.2f}",
        })
    # 人力工时明细加入"其他费用"分组
    labor_total_sum = 0.0
    for lh in data.get('labor_hours', []):
        labor_total_sum += float(lh.get('total', 0))
        fee_items.append({
            'item': lh.get('name', t('labor_hours', lang)),
            'spec': '',
            'category': t('labor_hours', lang),
            'unit_price': f'¥{float(lh.get("unit_price", 0)):.2f}',
            'quantity': str(lh.get('hours', 0)),
            'subtotal': f'¥{float(lh.get("total", 0)):.2f}',
        })
    fees_total += labor_total_sum
    if fee_items:
        module_groups.append({
            'name': t('other_fees', lang),
            'items': fee_items,
            'total': fees_total
        })

    buffer = BytesIO()

    class PDF(FPDF):
        def header(self):
            pass
        def footer(self):
            pass

    pdf = PDF()
    pdf.add_font('SimHei', '', FONT_REGULAR, uni=True)
    pdf.add_font('SimHei', 'B', FONT_BOLD, uni=True)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # 标题
    pdf.set_font('SimHei', 'B', 16)
    pdf.cell(0, 12, data.get('name', '报价单'), 0, 1, 'C')
    pdf.ln(5)

    # 基本信息
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('basic_info', lang), 0, 1)
    pdf.set_font('SimHei', '', 10)

    business_owner_name = ''
    if data.get('business_owner_id'):
        owner = User.query.get(data.get('business_owner_id'))
        if owner:
            business_owner_name = owner.real_name
    basic_info = [
        (t('project_name', lang), data.get('name', '')),
        (t('scheme_no', lang), data.get('scheme_no', '')),
        (t('business_owner', lang), business_owner_name),
        (t('version_no', lang), f'v{version_no}'),
        (t('currency', lang), currency)
    ]
    for key, value in basic_info:
        pdf.set_font('SimHei', 'B', 9)
        pdf.cell(40, 6, key, 1)
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 6, str(value), 1, 1)

    pdf.ln(5)

    # 物料及费用清单 (PDF)
    if module_groups:
        col_widths = [22, 38, 26, 20, 22, 12, 26, 24]  # 总约190
        headers = [t('module_col', lang), t('item_col', lang), t('spec_col', lang),
                   t('category_col', lang), t('unit_price_col', lang), t('qty_col', lang),
                   t('subtotal_col', lang), t('total_col', lang)]

        def draw_table_header():
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            for i, h in enumerate(headers):
                pdf.cell(col_widths[i], 6, h, 1, 0, 'C', True)
            pdf.ln()
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)

        def draw_group(y_start, group):
            item_count = len(group['items'])
            row_h = 5
            y_bottom = y_start + item_count * row_h

            # 使用PDF实际的左边距（默认10）
            x_start = pdf.l_margin if pdf.l_margin else 10
            x_module = x_start
            x_data = x_start + col_widths[0]
            x_total = x_start + col_widths[0] + sum(col_widths[1:7])
            x_end = x_start + sum(col_widths)

            # 模块列外边框
            pdf.rect(x_module, y_start, col_widths[0], item_count * row_h)
            # 合计列外边框
            pdf.rect(x_total, y_start, col_widths[7], item_count * row_h)

            # 数据区域边框
            pdf.line(x_data, y_start, x_data, y_bottom)
            pdf.line(x_total, y_start, x_total, y_bottom)

            # 列竖向分割线
            x_cur = x_data
            for j in range(1, 7):
                x_line = x_cur + col_widths[j]
                pdf.line(x_line, y_start, x_line, y_bottom)
                x_cur += col_widths[j]

            # 行水平分割线
            for i in range(1, item_count):
                y_line = y_start + i * row_h
                pdf.line(x_data, y_line, x_total, y_line)

            # 顶边和底边
            pdf.line(x_data, y_start, x_total, y_start)
            pdf.line(x_data, y_bottom, x_total, y_bottom)

            # 文字 - 模块列使用 multi_cell 自动换行
            text = group['name']
            line_height = 4
            lines_wrapped = pdf.multi_cell(col_widths[0], line_height, text, border=0, align='C', split_only=True)
            text_height = len(lines_wrapped) * line_height
            cell_height = item_count * row_h
            y_text = y_start + (cell_height - text_height) / 2
            pdf.set_xy(x_module, y_text)
            pdf.multi_cell(col_widths[0], line_height, text, 0, 'C')
            # 合计列
            pdf.set_xy(x_total, y_start + (item_count * row_h - 5) / 2)
            pdf.cell(col_widths[7], 5, f'¥{group["total"]:.2f}', 0, 0, 'C')

            # 数据单元格内容
            for i, item in enumerate(group['items']):
                y_row = y_start + i * row_h
                x_cur = x_data
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[1], row_h, str(item['item'])[:17], 0)
                x_cur += col_widths[1]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[2], row_h, str(item['spec'])[:12], 0)
                x_cur += col_widths[2]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[3], row_h, str(item['category'])[:9], 0)
                x_cur += col_widths[3]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[4], row_h, str(item['unit_price']), 0, 0, 'R')
                x_cur += col_widths[4]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[5], row_h, str(item['quantity']), 0, 0, 'C')
                x_cur += col_widths[5]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[6], row_h, str(item['subtotal']), 0, 0, 'R')

            # 移动到下一行
            pdf.set_y(y_bottom)
            return y_bottom

        def check_page_break(y_current, item_count, row_h=5):
            """检查是否需要换页，留出足够空间给当前模块"""
            try:
                y_current = float(y_current) if y_current is not None else 0
            except (ValueError, TypeError):
                y_current = 0
            pdf_h = float(pdf.h) if pdf.h else 297.0
            pdf_bm = float(pdf.b_margin) if pdf.b_margin else 10.0
            page_height = pdf_h - pdf_bm
            needed_height = item_count * row_h + 10  # 10为表头高度
            if y_current + needed_height > page_height:
                pdf.add_page()
                draw_table_header()
                new_y = pdf.get_y()
                return float(new_y) if new_y is not None else (pdf_h - pdf_bm - 10)
            return y_current

        # 开始绘制表格
        pdf.ln(5)
        pdf.set_font('SimHei', 'B', 11)
        pdf.cell(0, 8, t('material_list', lang), 0, 1)
        draw_table_header()

        y_start = pdf.get_y()
        for group in module_groups:
            if not group['items']:
                pdf.cell(sum(col_widths), 5, f'{group["name"]} - （无物料）', 1, 1, 'L')
                y_start = pdf.get_y()
                continue

            item_count = len(group['items'])
            # 检查是否需要换页
            y_start = check_page_break(y_start, item_count)

            y_end = draw_group(y_start, group)
            if y_end is None:
                y_end = y_start + len(group['items']) * 5
            y_start = y_end

        # 移动到下一页
        if y_start + 10 > (pdf.h if pdf.h else 297) - (pdf.b_margin if pdf.b_margin else 10):
            pdf.add_page()
    else:
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 5, '无物料及费用数据', 0, 1)

    pdf.ln(5)

    # 报价汇总
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('quote_summary', lang), 0, 1)

    pdf.set_font('SimHei', 'B', 9)
    pdf.set_fill_color(68, 114, 196)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(100, 6, t('project_col', lang), 1, 0, 'C', True)
    pdf.cell(0, 6, t('amount_col', lang), 1, 1, 'C', True)

    pdf.set_text_color(0, 0, 0)
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary_items = [
        (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
        (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
        (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
        (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
        (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
        (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
    ]
    pdf.set_font('SimHei', '', 9)
    for label, value in summary_items:
        pdf.cell(100, 5, label, 1)
        pdf.cell(0, 5, value, 1, 1, 'R')

    # 最终报价（高亮）
    pdf.set_font('SimHei', 'B', 10)
    pdf.set_fill_color(226, 239, 218)
    pdf.cell(100, 7, f"{t('grand_total', lang)}({currency})", 1, 0, 'C', True)
    pdf.cell(0, 7, f'{currency_symbol}{grand_total_converted:.2f}', 1, 1, 'R', True)

    # 输出
    pdf_bytes = pdf.output()
    with open(pdf_path, 'wb') as f:
        f.write(pdf_bytes)


# ==================== 快照数据读取 ====================


def get_version_snapshot_data(quotation_id, version_no):
    """获取版本快照数据并重建完整结构"""
    version = VersionSnapshot.query.filter_by(
        quotation_id=quotation_id,
        version_no=version_no
    ).first()
    if not version:
        return None

    import json
    data = json.loads(version.snapshot_data)

    # 重建设费用信息（快照已保存完整数据，仅做字段标准化）
    for fee in data.get('fees', []):
        fee['name'] = fee.get('fee_type', '') or fee.get('name', '')
        fee['amount'] = float(fee.get('amount', 0))
        fee['location'] = fee.get('position', 'internal')

    # 旧快照（unit_price=None）回退查询 Material（仅作为兜底）
    # 注意：保留 is_other 物料的 unit_price_override，不被 live 数据覆盖
    for module in data.get('modules', []):
        for mm in module.get('materials', []):
            mat = Material.query.get(mm.get('material_id'))
            if mat:
                mm['name'] = mat.name or ''
                mm['brand'] = mat.brand or ''
                mm['spec'] = mat.spec or ''
                # 只有非 is_other 物料才用 live 价；is_other 保留快照中的 override 价
                if not mm.get('is_other', False):
                    mm['unit_price'] = float(mat.unit_price) if mat.unit_price else 0
                mm['category'] = mat.category or 'standard'
                # 旧快照没有 unit_price_override，如果 brand='其他' 查 ModuleMaterial 表
                if mm.get('brand') == '其他' and mm.get('unit_price_override') is None:
                    mm_obj = ModuleMaterial.query.filter_by(
                        material_id=mm.get('material_id'),
                        module_id=module.get('id')
                    ).first()
                    if mm_obj and mm_obj.unit_price_override:
                        mm['unit_price_override'] = float(mm_obj.unit_price_override)

    return data


# ==================== 快照导出辅助 ====================


def _build_export_data_from_snapshot(snapshot_data_str):
    """从 snapshot_data 构建脱敏导出数据（用于高效渲染/导出）"""
    import json
    data = json.loads(snapshot_data_str) if isinstance(snapshot_data_str, str) else snapshot_data_str
    # 已在 normalize 后结构化为 {quotation, modules, fees, labor_hours}
    return data


# ==================== 版本合计计算 ====================


def calculate_version_totals(data):
    """计算版本数据的汇总"""
    modules = data.get('modules', [])
    fees = data.get('fees', [])

    # 分类统计物料
    category_totals = {'large': 0, 'standard': 0, 'other': 0}
    for module in modules:
        for mm in module.get('materials', []):
            is_other = mm.get('is_other', False)
            is_custom = mm.get('is_custom', False)
            if is_other:
                up = float(mm.get('unit_price_override', 0) or 0)
                if up == 0:
                    up = float(mm.get('unit_price', 0) or 0)
            elif is_custom:
                # migration 020: 自制件用 unit_price_override (= JSONB unit_price)
                up = float(mm.get('unit_price', 0) or 0)
                if up == 0:
                    up = float(mm.get('unit_price_override', 0) or 0)
            else:
                up = float(mm.get('unit_price', 0) or 0)
            amount = up * float(mm.get('quantity', 0) or 0)
            # 标准化 category：large/standard 保留，其他归入 other
            cat = mm.get('category', 'standard')
            if cat not in ('large', 'standard'):
                cat = 'other'
            category_totals[cat] = category_totals.get(cat, 0) + amount

    # 获取费用系数：优先用快照中的 coefficients，其次用报价单系数，最后用系统费率
    snapshot_coeff = data.get('coefficients', {}) or {}
    if snapshot_coeff:
        rate_large = snapshot_coeff.get('large', 1.0)
        rate_standard = snapshot_coeff.get('standard', 1.0)
        rate_other = snapshot_coeff.get('other', 1.0)
        fee_rates = snapshot_coeff
    elif data.get('quotation') and data['quotation'].get('coefficients'):
        fee_rates = data['quotation']['coefficients']
        rate_large = fee_rates.get('large', 1.0)
        rate_standard = fee_rates.get('standard', 1.0)
        rate_other = fee_rates.get('other', 1.0)
    else:
        fee_rates = {fr.category: float(fr.rate) for fr in FeeRate.query.all()}
        rate_large = fee_rates.get('large', 1.0)
        rate_standard = fee_rates.get('standard', 1.0)
        rate_other = fee_rates.get('other', 1.0)

    material_with_rates = (
        category_totals.get('large', 0) * rate_large +
        category_totals.get('standard', 0) * rate_standard +
        category_totals.get('other', 0) * rate_other
    )

    material_total = sum(category_totals.values())
    fees_total = sum(float(fee.get('amount', 0)) for fee in fees)
    # 运输包装 + 差旅人天 + 差旅人次 小计
    packing_total = sum(float(e.get('total') or 0) for e in data.get('packing_entries', []))
    person_days_total = sum(float(e.get('total') or 0) for e in data.get('person_days_entries', []))
    person_trips_total = sum(float(e.get('total') or 0) for e in data.get('person_trip_entries', []))
    fees_total += packing_total + person_days_total + person_trips_total
    labor_total = sum(float(l.get('total', 0)) for l in data.get('labor_hours', []))
    subtotal = material_with_rates + fees_total + labor_total

    # 获取对外利润率（从报价单或根级获取）
    profit_rate = 0.0
    if data.get('profit_rate') is not None:
        profit_rate = float(data.get('profit_rate', 0))
    elif data.get('quotation', {}).get('profit_rate') is not None:
        profit_rate = float(data['quotation'].get('profit_rate', 0))
    profit_amount = subtotal * profit_rate
    subtotal_with_profit = subtotal + profit_amount

    # 获取税率（从报价单或根级获取）
    tax_rate = 0.13
    if data.get('tax_rate') is not None:
        tax_rate = float(data.get('tax_rate', 0.13))
    elif data.get('quotation', {}).get('tax_rate') is not None:
        tax_rate = float(data['quotation'].get('tax_rate', 0.13))
    grand_total = subtotal_with_profit * (1 + tax_rate)
    tax_amount = grand_total - subtotal_with_profit

    return {
        'material_total': material_total,
        'material_total_with_rates': material_with_rates,
        'fees_total': fees_total,
        'labor_total': labor_total,
        'subtotal': subtotal,
        'profit_rate': profit_rate,
        'profit_amount': profit_amount,
        'subtotal_with_profit': subtotal_with_profit,
        'tax_rate': tax_rate,
        'tax_amount': tax_amount,
        'grand_total': grand_total,
        'fee_rates': fee_rates
    }


# ==================== PDF 渲染（共享） ====================


def render_pdf_bytes(tables, currency, currency_symbol, material_with_rates, table2_total, table3_total,
                     profit_rate, subtotal_with_profit, tax_rate, tax_amount, grand_total_converted,
                     quotation, data, lang):
    """共享 PDF 渲染逻辑，中英文通用"""

    class PDF(FPDF):
        def header(self): pass
        def footer(self): pass

    pdf = PDF()
    pdf.add_font('SimHei', '', FONT_REGULAR, uni=True)
    pdf.add_font('SimHei', 'B', FONT_BOLD, uni=True)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # 标题
    pdf.set_font('SimHei', 'B', 16)
    pdf.cell(0, 12, t('quote', lang), 0, 1, 'C')
    pdf.ln(5)

    # 基本信息
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('basic_info', lang), 0, 1)
    pdf.set_font('SimHei', '', 10)

    # 从 quotation 或 data 中获取基本信息
    q_name = quotation.name if quotation else data.get('name', '')
    q_scheme_no = quotation.scheme_no if quotation else data.get('scheme_no', '')
    q_owner = quotation.business_owner.real_name if quotation and quotation.business_owner else ''
    q_created = quotation.created_at.strftime('%Y-%m-%d %H:%M') if quotation and quotation.created_at else ''
    q_version = data.get('version_no', '') if data else ''

    basic_info = [
        (t('project_name', lang), q_name or ''),
        (t('scheme_no', lang), q_scheme_no or ''),
        (t('business_owner', lang), q_owner),
        (t('created_at', lang), q_created),
        (t('currency', lang), currency),
    ]
    if q_version:
        basic_info.append((t('version_no', lang), f'v{q_version}'))

    for key, value in basic_info:
        pdf.set_font('SimHei', 'B', 9)
        pdf.cell(40, 6, key, 1)
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 6, str(value), 1, 1)

    pdf.ln(5)

    # ========== 合并表格 ==========
    if tables['table1']['rows'] or tables['table2']['rows'] or tables['table3']['rows']:
        unified_col_widths = [60, 35, 30, 20, 45]

        # 1. 设备硬件
        if tables['table1']['rows']:
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(sum(unified_col_widths), 5, '【设备硬件】', 1, 1, 'L', True)
            pdf.set_fill_color(217, 217, 217)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(unified_col_widths[0], 6, t('module_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[1], 6, t('brand_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[2], 6, t('qty_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[3], 6, t('unit_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, t('total_col', lang), 1, 1, 'C', True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
            for row in tables['table1']['rows']:
                pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
                pdf.cell(unified_col_widths[0], 5, str(row['module_name'])[:20], 1)
                pdf.cell(unified_col_widths[1], 5, str(row['brand'])[:10], 1)
                qty_val = row['quantity']
                pdf.cell(unified_col_widths[2], 5, str(int(qty_val)) if qty_val == int(qty_val) else str(qty_val), 1, 0, 'C')
                pdf.cell(unified_col_widths[3], 5, str(row['unit']), 1, 0, 'C')
                pdf.cell(unified_col_widths[4], 5, f'¥{row["subtotal"]:.2f}', 1, 1, 'R')
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(226, 239, 218)
            pdf.cell(sum(unified_col_widths[:4]), 6, t('subtotal_label', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, f'¥{tables["table1"]["total"]:.2f}', 1, 1, 'R', True)

        # 2. 设备人力
        if tables['table2']['rows']:
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(sum(unified_col_widths), 5, '【设备人力】', 1, 1, 'L', True)
            pdf.set_fill_color(217, 217, 217)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(unified_col_widths[0] + unified_col_widths[1], 6, t('division_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[2], 6, t('qty_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[3], 6, t('unit_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, t('total_col', lang), 1, 1, 'C', True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
            for row in tables['table2']['rows']:
                pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
                pdf.cell(unified_col_widths[0] + unified_col_widths[1], 5, str(row['division'])[:20], 1)
                pdf.cell(unified_col_widths[2], 5, f'{row["hours"]:.1f}', 1, 0, 'C')
                pdf.cell(unified_col_widths[3], 5, 'H', 1, 0, 'C')
                pdf.cell(unified_col_widths[4], 5, f'¥{row["subtotal"]:.2f}', 1, 1, 'R')
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(226, 239, 218)
            pdf.cell(sum(unified_col_widths[:4]), 6, t('subtotal_label', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, f'¥{tables["table2"]["total"]:.2f}', 1, 1, 'R', True)

        # 3. 其他合计
        if tables['table3']['rows']:
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(sum(unified_col_widths), 5, '【其他】', 1, 1, 'L', True)
            pdf.set_fill_color(217, 217, 217)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.cell(unified_col_widths[0] + unified_col_widths[1], 6, t('fee_name_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[2], 6, t('qty_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[3], 6, t('unit_col', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, t('unit_price_col', lang), 1, 1, 'C', True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
            for row in tables['table3']['rows']:
                pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
                pdf.cell(unified_col_widths[0] + unified_col_widths[1], 5, str(row['name'])[:20], 1)
                qty_val = row.get('quantity', 1)
                pdf.cell(unified_col_widths[2], 5, str(int(qty_val)) if isinstance(qty_val, (int, float)) and qty_val == int(qty_val) else str(qty_val), 1, 0, 'C')
                pdf.cell(unified_col_widths[3], 5, str(row.get('unit', 'SET')), 1, 0, 'C')
                pdf.cell(unified_col_widths[4], 5, f'¥{row.get("unit_price", 0):.2f}', 1, 1, 'R')
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(226, 239, 218)
            pdf.cell(sum(unified_col_widths[:4]), 6, t('subtotal_label', lang), 1, 0, 'C', True)
            pdf.cell(unified_col_widths[4], 6, f'¥{tables["fees_subtotal"]:.2f}', 1, 1, 'R', True)

        # 最终报价
        pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
        pdf.set_font('SimHei', 'B', 9)
        pdf.set_fill_color(255, 255, 204)
        pdf.cell(sum(unified_col_widths[:2]), 6, t('grand_total', lang), 1, 0, 'C', True)
        pdf.cell(sum(unified_col_widths[2:4]), 6, currency, 1, 0, 'C', True)
        pdf.cell(unified_col_widths[4], 6, f'{currency_symbol}{grand_total_converted:.2f}', 1, 1, 'R', True)

    return pdf.output()


# ==================== PDF 表格构建（快照版）====================


def _build_pdf_tables_from_snapshot(data, quotation, coeff, profit_rate_override=None, tax_rate_override=None):
    """从快照数据构建 PDF tables（与 _build_pdf_tables 逻辑一致）"""
    rate_large = float(coeff.get('large', 1.0))
    rate_standard = float(coeff.get('standard', 1.0))
    rate_other = float(coeff.get('other', 1.0))

    modules_data = data.get('modules', [])
    fees_data = data.get('fees', [])
    labor_data = data.get('labor_hours', [])
    packing_data = data.get('packing_entries', [])
    person_days_data = data.get('person_days_entries', [])
    person_trip_data = data.get('person_trip_entries', [])

    # table1 设备硬件
    # 优先用快照中的 quotation type，否则用 live 数据
    quotation_type = data.get('quotation', {}).get('type') or (quotation.type if quotation else 'standalone')
    if quotation_type == 'line':
        from collections import defaultdict
        modules_by_quotation = defaultdict(list)
        for mod in modules_data:
            modules_by_quotation[mod.get('quotation_id', quotation.id if quotation else None)].append(mod)
        child_names = {c.id: c.name for c in quotation.children.all()} if quotation and hasattr(quotation, 'children') else {}
        table1_rows = []
        table1_total = 0.0
        for qid, mod_list in modules_by_quotation.items():
            group_total = 0.0
            for mod in mod_list:
                for mm in mod.get('materials', []):
                    cat = mm.get('category', 'standard')
                    rate = rate_large if cat == 'large' else (rate_standard if cat == 'standard' else rate_other)
                    # is_other 物料：用 unit_price_override
                    is_other = mm.get('is_other', False)
                    if is_other:
                        up = float(mm.get('unit_price_override', 0) or 0)
                        if up == 0:
                            up = float(mm.get('unit_price', 0) or 0)
                    else:
                        up = float(mm.get('unit_price', 0) or 0)
                    group_total += up * rate * float(mm.get('quantity', 0) or 0)
            display_name = '(线体)' if qid == (quotation.id if quotation else None) else child_names.get(qid, f'子项目{qid}')
            table1_rows.append({'module_name': display_name, 'brand': 'RS', 'quantity': 1, 'unit': 'SET', 'subtotal': group_total})
            table1_total += group_total
    else:
        # 单机/独立报价单：每模块一行（模块内所有物料×系数合计）
        table1_rows = []
        table1_total = 0.0
        for mod in modules_data:
            mod_qid = mod.get('quotation_id', quotation.id)
            if mod_qid != quotation.id:
                continue
            module_amount = 0.0
            non_other_materials = []
            for mm in mod.get('materials', []):
                # 推断 is_other：快照有 is_other 字段用字段；否则用 brand='其他' 或 category='other' 推断
                is_other = mm.get('is_other', False)
                if not is_other:
                    brand = mm.get('brand', '')
                    cat = mm.get('category', '')
                    is_other = (brand == '其他' or cat == 'other' or cat == '其他件')

                if is_other:
                    # 其他件：用 unit_price_override（快照已保存）
                    up = float(mm.get('unit_price_override', 0) or 0)
                    if up == 0:
                        up = float(mm.get('unit_price', 0) or 0)
                    qty = float(mm.get('quantity', 0) or 0)
                    cat = 'other'
                else:
                    up = float(mm.get('unit_price', 0) or 0)
                    qty = float(mm.get('quantity', 0) or 0)
                    cat = mm.get('category', 'standard')
                    if cat not in ('large', 'standard'):
                        cat = 'other'
                    if up > 0:
                        non_other_materials.append(mm)

                rate = (rate_large if cat == 'large'
                        else (rate_standard if cat == 'standard' else rate_other))
                module_amount += up * qty * rate

            if module_amount > 0:
                # 品牌逻辑：只有其他→其他；1个非其他→该物料品牌；多个非其他→RS
                if len(non_other_materials) == 0:
                    brand = '其他'
                elif len(non_other_materials) == 1:
                    brand = non_other_materials[0].get('brand', '') or '其他'
                else:
                    brand = 'RS'

                table1_rows.append({
                    'module_name': mod.get('name', ''),
                    'brand': brand,
                    'quantity': 1,
                    'unit': 'SET',
                    'subtotal': module_amount
                })
                table1_total += module_amount

    # table2 设备人力
    table2_rows = []
    table2_total = 0.0
    for lh in labor_data:
        table2_rows.append({'division': lh.get('name', ''), 'hours': float(lh.get('hours', 0)), 'subtotal': float(lh.get('total', 0))})
        table2_total += float(lh.get('total', 0))

    # ===== 三大新费用汇总（与 Tab 版 _build_pdf_tables 完全一致）=====
    total_person_days = 0.0
    total_person_days_count = 0.0
    for entry in person_days_data:
        up = float(entry.get('unit_price', 0) or 0)
        days = float(entry.get('person_days') or entry.get('days') or entry.get('quantity') or 0)
        total_person_days += up * days
        total_person_days_count += days

    total_person_trips = 0.0
    total_person_trips_count = 0.0
    for entry in person_trip_data:
        up = float(entry.get('unit_price', 0) or 0)
        vf = float(entry.get('visa_fee', 0) or 0)
        # 已存 total = quantity * (unit_price + visa_fee) ，即与 Tab 版 cat_code != 'domestic' 一致
        total = float(entry.get('total', 0) or 0)
        if total == 0:
            # 兼容没有 total 字段的旧快照
            cat_code = entry.get('travel_category_code', '') or entry.get('cat_code', '')
            count = float(entry.get('person_count') or entry.get('quantity') or 0)
            total = count * (up + (vf if cat_code != 'domestic' else 0))
        else:
            count = float(entry.get('person_count') or entry.get('quantity') or 0)
        total_person_trips += total
        total_person_trips_count += count

    total_packing = 0.0
    total_packing_count = 0.0
    for entry in packing_data:
        up = float(entry.get('unit_price', 0) or 0)
        qty = float(entry.get('quantity', 0) or 0)
        total_packing += up * qty
        total_packing_count += qty

    # ===== 表格3：其他合计（与 Tab 版完全一致）=====
    table3_rows = []
    table3_total = 0.0

    if total_person_days > 0:
        table3_rows.append({
            'name': '差旅住宿费',
            'quantity': round(total_person_days_count, 1),
            'unit': '人天',
            'unit_price': round(total_person_days, 2),
        })
        table3_total += total_person_days

    if total_person_trips > 0:
        table3_rows.append({
            'name': '差旅交通签证费',
            'quantity': round(total_person_trips_count, 0),
            'unit': '人次',
            'unit_price': round(total_person_trips, 2),
        })
        table3_total += total_person_trips

    if total_packing > 0:
        table3_rows.append({
            'name': '设备包装运输费',
            'quantity': round(total_packing_count, 0),
            'unit': '单元',
            'unit_price': round(total_packing, 2),
        })
        table3_total += total_packing

    # 动态费用（OtherFee 记录：每条一行）
    dynamic_fees_total = 0.0
    for fee in fees_data:
        fee_amount = float(fee.get('amount', 0) or 0)
        table3_rows.append({
            'name': fee.get('fee_type', '') or fee.get('name', ''),
            'quantity': 1,
            'unit': 'SET',
            'unit_price': round(fee_amount, 2),
        })
        dynamic_fees_total += fee_amount
    table3_total += dynamic_fees_total

    # 项目利润和项目税额（基于 base table3_total）
    base_table3_total = table3_total
    profit_rate_val = profit_rate_override if profit_rate_override is not None else float(data.get('profit_rate', 0) or 0)
    tax_rate_val = tax_rate_override if tax_rate_override is not None else float(data.get('tax_rate', 0) or 0)
    profit_amount = (table1_total + table2_total + base_table3_total) * profit_rate_val
    subtotal_with_profit = table1_total + table2_total + base_table3_total + profit_amount
    tax_amount = subtotal_with_profit * tax_rate_val

    # 项目利润行
    table3_rows.append({
        'name': '项目利润',
        'quantity': 1,
        'unit': 'SET',
        'unit_price': round(profit_amount, 2),
    })
    # 项目税额行
    table3_rows.append({
        'name': '项目税额',
        'quantity': 1,
        'unit': 'SET',
        'unit_price': round(tax_amount, 2),
    })
    table3_total += profit_amount
    table3_total += tax_amount

    # fees_subtotal = 基础费用 + 利润 + 税额（与 Tab 版一致）
    fees_subtotal = table3_total

    return {
        'table1': {'rows': table1_rows, 'total': round(table1_total, 2)},
        'table2': {'rows': table2_rows, 'total': round(table2_total, 2)},
        'table3': {'rows': table3_rows, 'total': round(table3_total, 2)},
        'fees_subtotal': fees_subtotal,
        'profit_amount': profit_amount,
        'tax_amount': tax_amount,
        'subtotal_with_profit': subtotal_with_profit,
    }


# ==================== PDF 汇总计算 ====================


def calculate_totals(tables, coeff, currency):
    """计算 PDF 汇总数据
    tables 已含 profit_amount / tax_amount / subtotal_with_profit
    （_build_pdf_tables 和 _build_pdf_tables_from_snapshot 都算好了）
    """
    material_with_rates = tables['table1']['total']
    table2_total = tables['table2']['total']
    fees_subtotal = tables['fees_subtotal']
    # base_subtotal = 1+2+table3 基础部分（不含利润/税）
    profit_amount = tables.get('profit_amount', 0)
    tax_amount = tables.get('tax_amount', 0)
    subtotal_with_profit = tables.get('subtotal_with_profit', material_with_rates + table2_total + fees_subtotal - profit_amount - tax_amount)
    # 从 subtotal_with_profit 反推 profit_rate（如果需要展示）
    base = subtotal_with_profit - profit_amount
    profit_rate = (profit_amount / base) if base else 0
    # 反推 tax_rate
    tax_rate = (tax_amount / subtotal_with_profit) if subtotal_with_profit else 0
    grand_total = subtotal_with_profit + tax_amount
    subtotal = material_with_rates + table2_total + fees_subtotal
    return {
        'material_with_rates': material_with_rates,
        'table2_total': table2_total,
        'fees_subtotal': fees_subtotal,
        'subtotal': subtotal,
        'profit_rate': profit_rate,
        'subtotal_with_profit': subtotal_with_profit,
        'tax_rate': tax_rate,
        'tax_amount': tax_amount,
        'grand_total': grand_total,
    }


# ==================== 版本 PDF 生成（归档时调用）====================


def generate_version_pdfs(quotation_id, version_no):
    """归档时生成中英两个 PDF，保存路径到 version.pdf_file，与 export_pdf 逻辑完全一致"""
    import json, os
    print(f"[gen_pdfs] qid={quotation_id} vno={version_no} start", flush=True)
    version = VersionSnapshot.query.filter_by(quotation_id=quotation_id, version_no=version_no).first()
    if not version:
        print(f"[gen_pdfs] version not found", flush=True)
        return

    data = get_version_snapshot_data(quotation_id, version_no)
    if not data:
        print(f"[gen_pdfs] get_version_snapshot_data returned None", flush=True)
        return
    print(f"[gen_pdfs] data loaded: type={data.get('type')} modules={len(data.get('modules',[]))}", flush=True)

    quotation = Quotation.query.get(quotation_id)
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    coeff = (quotation.coefficients if hasattr(quotation, 'coefficients') and quotation.coefficients else
             {fr.category: float(fr.rate) for fr in FeeRate.query.all()})

    # 重建 tables 数据（与 _build_pdf_tables 逻辑一致）
    tables = _build_pdf_tables_from_snapshot(data, quotation, coeff)

    # 直接用 tables 中已含利润/税额的正确汇总值
    totals = {
        'material_with_rates': tables['table1']['total'],
        'table2_total': tables['table2']['total'],
        'fees_subtotal': tables['fees_subtotal'],
        'subtotal': tables['table1']['total'] + tables['table2']['total'] + tables['fees_subtotal'],
        'profit_amount': tables.get('profit_amount', 0),
        'subtotal_with_profit': tables.get('subtotal_with_profit', 0),
        'tax_amount': tables.get('tax_amount', 0),
        'tax_rate': float(quotation.tax_rate) if quotation and quotation.tax_rate else 0,
        'grand_total': tables.get('subtotal_with_profit', 0) + tables.get('tax_amount', 0),
    }
    # 利润率
    total_before_profit = tables['table1']['total'] + tables['table2']['total'] + tables['fees_subtotal'] - tables.get('profit_amount', 0) - tables.get('tax_amount', 0)
    totals['profit_rate'] = totals['profit_amount'] / total_before_profit if total_before_profit else 0
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)
    # 保存到 static/versions/（统一存放在 backend_fastapi/static/versions/）
    # __file__ = core/services/export_service.py → 3 次 dirname 升到 backend_fastapi
    base_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'static', 'versions')
    os.makedirs(base_dir, exist_ok=True)

    pdf_paths = {}
    word_path = None
    for lang in ['zh', 'en']:
        pdf_bytes = render_pdf_bytes(
            tables, currency, currency_symbol,
            totals['material_with_rates'], totals['table2_total'], totals['fees_subtotal'],
            totals['profit_rate'], totals['subtotal_with_profit'], totals['tax_rate'],
            totals['tax_amount'], grand_total_converted,
            quotation, data, lang
        )
        file_path = os.path.join(base_dir, f'quotation_{quotation_id}_v{version_no}_{lang}.pdf')
        with open(file_path, 'wb') as f:
            f.write(pdf_bytes)
        pdf_paths[lang] = file_path

    # Word 导出（中文版）— 与 Tab 导出的 Word 格式对齐
    try:
        word_path = os.path.join(base_dir, f'quotation_{quotation_id}_v{version_no}.docx')
        from core.services.export_service import _generate_version_word, calculate_version_totals
        # 复用 _generate_version_word 所需格式（需要 totals 含 grand_total 等）
        # 将 _build_pdf_tables_from_snapshot 的结果转为 calculate_version_totals 期望的格式
        # 但这里直接用 totals 已有的字段构造
        word_totals = {
            'material_total_with_rates': totals['material_with_rates'],
            'labor_total': totals['table2_total'],
            'fees_total': totals['fees_subtotal'] - totals.get('profit_amount', 0) - totals.get('tax_amount', 0),
            'fee_rates': coeff,
            'profit_rate': totals['profit_rate'],
            'subtotal_with_profit': totals['subtotal_with_profit'],
            'tax_rate': totals['tax_rate'],
            'tax_amount': totals['tax_amount'],
            'grand_total': totals['subtotal_with_profit'] + totals['tax_amount'],
        }
        _generate_version_word(word_path, quotation_id, version_no, data, word_totals, quotation, 'zh')
    except Exception as _we:
        import traceback
        print(f"生成 Word 失败: {_we}", flush=True)
        traceback.print_exc()

    version.pdf_file = json.dumps(pdf_paths)
    version.word_file = word_path

    # 提交到数据库（持久化 pdf_file 字段）
    try:
        from db import db as _db
        _db.session.add(version)
        _db.session.commit()
    except Exception as _e:
        try:
            _db.session.rollback()
        except Exception:
            pass
        print(f"保存 version.pdf_file 失败: {_e}")
