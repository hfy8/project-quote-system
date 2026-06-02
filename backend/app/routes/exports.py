"""导出功能模块 - 支持 Word/Excel/PDF 格式"""
from flask import Blueprint, send_file, jsonify, request
from io import BytesIO
from datetime import datetime

# Word 导出
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Excel 导出
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# PDF 导出 - 使用 fpdf2 支持中文
from fpdf import FPDF

# 中文字体路径（相对路径，字体跟随项目）
import os
FONT_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'fonts')
FONT_REGULAR = os.path.join(FONT_DIR, 'simhei.ttf')  # 黑体（支持中文）
FONT_BOLD = os.path.join(FONT_DIR, 'simhei.ttf')       # 黑体（暂无粗体，用黑体代替）

# 中英文标签映射
I18N = {
    'zh': {
        'quote': '报价单', 'basic_info': '基本信息', 'project_name': '项目名称',
        'scheme_no': '方案编号', 'business_owner': '业务负责人', 'created_at': '创建时间',
        'currency': '报价币种', 'material_list': '物料及费用清单',
        'module_col': '模块', 'item_col': '项目', 'spec_col': '规格',
        'category_col': '分类/品牌', 'unit_price_col': '单价', 'qty_col': '数量',
        'subtotal_col': '小计', 'total_col': '合计',
        'quote_summary': '报价汇总', 'project_col': '项目', 'amount_col': '金额',
        'material_total': '物料合计', 'fee_total': '费用合计', 'labor_total': '人力合计',
        'profit_rate': '利润率', 'subtotal_with_profit': '含利润率小计',
        'tax_rate': '税率', 'tax_amount': '税额', 'grand_total': '最终报价',
        'other_fees': '其他费用', 'unnamed_module': '未命名模块',
        'quote_summary': '报价汇总',
        'internal': '厂内', 'external': '厂外',
        'no_materials': '（无物料）', 'no_data': '无物料及费用数据',
        'labor_hours': '人力工时', 'version': '版本', 'version_no': '版本号',
        'profit_rate_label': '利润率', 'subtotal_label': '含利润率小计',
        'unit_h': 'h', 'tax_included': '含税报价',
    },
    'en': {
        'quote': 'QUOTATION', 'basic_info': 'Basic Information', 'project_name': 'Project Name',
        'scheme_no': 'Scheme No.', 'business_owner': 'Sales Owner', 'created_at': 'Created At',
        'currency': 'Currency', 'material_list': 'Material & Cost List',
        'module_col': 'Module', 'item_col': 'Item', 'spec_col': 'Spec',
        'category_col': 'Category/Brand', 'unit_price_col': 'Unit Price', 'qty_col': 'Qty',
        'subtotal_col': 'Subtotal', 'total_col': 'Total',
        'quote_summary': 'Quote Summary', 'project_col': 'Item', 'amount_col': 'Amount',
        'material_total': 'Material Total', 'fee_total': 'Fee Total', 'labor_total': 'Labor Total',
        'profit_rate': 'Profit Rate', 'subtotal_with_profit': 'Subtotal w/ Profit',
        'tax_rate': 'Tax Rate', 'tax_amount': 'Tax Amount', 'grand_total': 'Grand Total',
        'other_fees': 'Other Fees', 'unnamed_module': 'Unnamed Module',
        'quote_summary': 'Quote Summary',
        'internal': 'Internal', 'external': 'External',
        'no_materials': '(No Materials)', 'no_data': 'No material or cost data',
        'labor_hours': 'Labor Hours', 'version': 'Version', 'version_no': 'Version No.',
        'profit_rate_label': 'Profit Rate', 'subtotal_label': 'Subtotal w/ Profit',
        'unit_h': 'h', 'tax_included': 'Tax Included',
    }
}

def t(key, lang='zh'):
    return I18N.get(lang, I18N['zh']).get(key, key)


def get_name(obj, lang='zh'):
    """获取对象的本地化名称，英文时优先用 name_en"""
    if lang == 'en':
        en = getattr(obj, 'name_en', None)
        if en:
            return en
    return obj.name


def get_fee_type_name(fee_type_id, lang='zh'):
    """根据 fee_type ID 或名称字符串获取费用类型名称，英文时优先用 name_en"""
    if not fee_type_id:
        return fee_type_id or ''
    # 如果是整数ID，按ID查
    if isinstance(fee_type_id, int):
        ft = db.session.get(FeeType, fee_type_id)
        if ft:
            return ft.name_en if (lang == 'en' and ft.name_en) else ft.name
        return str(fee_type_id)
    # 如果是字符串（可能是中文名或英文名），按名称查
    ft = FeeType.query.filter(FeeType.name == str(fee_type_id)).first()
    if ft:
        return ft.name_en if (lang == 'en' and ft.name_en) else ft.name
    # 也尝试查 name_en
    if lang == 'en':
        ft = FeeType.query.filter(FeeType.name_en == str(fee_type_id)).first()
        if ft:
            return ft.name_en
    return str(fee_type_id)

from app.models.quotation import Quotation
from app.models.module import Module
from app.models.material import ModuleMaterial, Material
from app.models.fee import OtherFee, FeeType
from app.models.fee_rate import FeeRate
from app.models.exchange_rate import ExchangeRate
from app.models.version import VersionSnapshot
from app.models.user import User
from app.models.labor_hour import LaborHour

export_bp = Blueprint('exports', __name__)


def get_quotation_with_details(quotation_id):
    """获取报价单及其详细信息"""
    quotation = Quotation.query.get(quotation_id)
    if not quotation:
        return None
    
    # 获取模块及物料
    modules = Module.query.filter_by(quotation_id=quotation_id).all()
    for module in modules:
        module.materials = ModuleMaterial.query.filter_by(module_id=module.id).all()
    
    # 获取其他费用
    fees = OtherFee.query.filter_by(quotation_id=quotation_id).all()
    labor_hours = LaborHour.query.filter_by(quotation_id=quotation_id).all()

    return quotation, modules, fees, labor_hours


def calculate_totals_with_rates(quotation, modules, fees, labor_hours):
    """计算汇总数据（含费用系数和税率）"""
    # 分类统计物料
    category_totals = {'large': 0, 'standard': 0, 'other': 0}
    for module in modules:
        for mm in module.materials:
            if mm.material:
                cat = mm.material.category or 'standard'
                amount = float(mm.material.unit_price or 0) * mm.quantity
                category_totals[cat] = category_totals.get(cat, 0) + amount
    
    # 获取费用系数（优先使用报价单私有系数，否则用全局配置）
    coeff = (quotation.coefficients if hasattr(quotation, 'coefficients') and quotation.coefficients else
             {fr.category: float(fr.rate) for fr in FeeRate.query.all()})
    rate_large = float(coeff.get('large', 1.0))
    rate_standard = float(coeff.get('standard', 1.0))
    rate_other = float(coeff.get('other', 1.0))
    fee_rates = coeff
    
    # 计算物料含系数小计
    material_with_rates = (
        category_totals.get('large', 0) * rate_large +
        category_totals.get('standard', 0) * rate_standard +
        category_totals.get('other', 0) * rate_other
    )
    
    material_total = sum(category_totals.values())
    fees_total = sum(float(fee.amount or 0) for fee in fees)
    labor_total = sum(float(l.total or 0) for l in labor_hours)
    subtotal = material_with_rates + fees_total + labor_total
    
    # 对外利润率
    profit_rate = float(quotation.profit_rate) if quotation.profit_rate else 0.0
    subtotal_with_profit = subtotal * (1 + profit_rate)
    
    # 获取税率
    tax_rate = float(quotation.tax_rate) if quotation.tax_rate else 0.13
    tax_amount = subtotal_with_profit * tax_rate
    grand_total = subtotal_with_profit + tax_amount
    
    return {
        'material_total': material_total,
        'material_total_with_rates': material_with_rates,
        'fees_total': fees_total,
        'labor_total': labor_total,
        'subtotal': subtotal,
        'profit_rate': profit_rate,
        'subtotal_with_profit': subtotal_with_profit,
        'tax_rate': tax_rate,
        'tax_amount': tax_amount,
        'grand_total': grand_total,
        'category_totals': category_totals,
        'fee_rates': fee_rates
    }


def convert_currency(amount, to_currency, exchange_rates):
    """汇率转换"""
    if to_currency == 'CNY':
        return amount
    # rate 是 "1 目标货币 = X CNY"
    rate = exchange_rates.get(to_currency, 1)
    return amount / rate


def get_currency_symbol(currency):
    """获取币种符号"""
    symbols = {'CNY': '¥', 'USD': '$', 'EUR': '€', 'GBP': '£', 'JPY': '¥'}
    return symbols.get(currency, '¥')


def generate_version_files(quotation_id, version_no, snapshot_data, data=None, lang='zh'):
    """为版本生成Word和PDF文件，返回文件路径字典"""
    import os
    import json
    
    # 创建版本文件存储目录
    base_dir = f'/mnt/c/Users/rs8568/Desktop/Project/project-quote-system/backend/versions/{quotation_id}'
    os.makedirs(base_dir, exist_ok=True)
    
    word_path = os.path.join(base_dir, f'v{version_no}.docx')
    pdf_zh_path = os.path.join(base_dir, f'v{version_no}_zh.pdf')
    pdf_en_path = os.path.join(base_dir, f'v{version_no}_en.pdf')
    
    # 获取报价单信息（包含币种）
    quotation = Quotation.query.get(quotation_id)
    
    if data is None:
        data = json.loads(snapshot_data) if isinstance(snapshot_data, str) else snapshot_data
    
    # 填充物料详细信息（与 get_version_snapshot_data 相同逻辑）
    for module in data.get('modules', []):
        for mm in module.get('materials', []):
            mat = Material.query.get(mm.get('material_id'))
            if mat:
                mm['name'] = mat.name or ''
                mm['brand'] = mat.brand or ''
                mm['spec'] = mat.spec or ''
                mm['unit_price'] = float(mat.unit_price or 0)
                mm['category'] = mat.category or 'standard'
    
    # 计算汇总
    totals = calculate_version_totals(data)
    
    # 生成 Word 文件
    _generate_version_word(word_path, quotation_id, version_no, data, totals, quotation, 'zh')
    
    # 生成 PDF 文件（中英文各一份）
    _generate_version_pdf(pdf_zh_path, quotation_id, version_no, data, totals, quotation, 'zh')
    _generate_version_pdf(pdf_en_path, quotation_id, version_no, data, totals, quotation, 'en')
    
    import json
    pdf_file_json = json.dumps({'zh': pdf_zh_path, 'en': pdf_en_path})
    return {'word': word_path, 'pdf': pdf_file_json}


def _generate_version_word(word_path, quotation_id, version_no, data, totals, quotation=None, lang='zh'):
    """生成版本 Word 文件"""
    doc = Document()
    title = doc.add_heading(f'{data.get("name", t("quote", lang))}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 币种和汇率
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)
    
    # 基本信息
    doc.add_heading(t('basic_info', lang), level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    business_owner_name = ''
    if data.get('business_owner_id'):
        owner = User.query.get(data.get('business_owner_id'))
        if owner:
            business_owner_name = owner.real_name
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    info_data = [
        (t('project_name', lang), data.get('name', '')),
        (t('scheme_no', lang), data.get('scheme_no', '')),
        (t('business_owner', lang), business_owner_name),
        (t('version_no', lang), f'v{version_no}'),
        (t('currency', lang), currency)
    ]
    for i, (key, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # 整合表格
    doc.add_heading(t('material_list', lang), level=1)
    
    module_groups = []
    rate_large = totals.get('fee_rates', {}).get('large', 1.0)
    rate_standard = totals.get('fee_rates', {}).get('standard', 1.0)
    rate_other = totals.get('fee_rates', {}).get('other', 1.0)
    for module in data.get('modules', []):
        items = []
        module_total = 0
        for mm in module.get('materials', []):
            cat = mm.get('category', 'standard')
            rate = rate_large if cat == 'large' else (rate_standard if cat == 'standard' else rate_other)
            unit_price_with_rate = float(mm.get('unit_price', 0) or 0) * rate
            item_total = unit_price_with_rate * float(mm.get('quantity', 0) or 0)
            module_total += item_total
            items.append({
                'item': mm.get('name', ''),
                'spec': mm.get('spec', ''),
                'category': mm.get('brand', ''),
                'unit_price': f'¥{unit_price_with_rate:.2f}',
                'quantity': str(mm.get('quantity', 0)),
                'subtotal': f'¥{item_total:.2f}',
            })
        module_groups.append({
            'name': module.get('name', t('unnamed_module', lang)),
            'items': items,
            'total': module_total
        })
    
    fees_total = sum(float(f.get('amount', 0)) for f in data.get('fees', []))
    fee_items = []
    for fee in data.get('fees', []):
        location_text = t('internal', lang) if fee.get('location') == 'internal' else (t('external', lang) if fee.get('location') == 'external' else '')
        fee_items.append({
            'item': get_fee_type_name(fee.get('fee_type', ''), lang) or fee.get('name', ''),
            'spec': location_text,
            'category': get_fee_type_name(fee.get('fee_type', ''), lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f'¥{float(fee.get("amount", 0)):.2f}',
        })
    # 人力工时明细加入"其他费用"分组
    labor_total_sum = 0.0
    for lh in data.get('labor_hours', []):
        labor_total_sum += float(lh.get('total', 0))
        fee_items.append({
            'item': lh.get('name', '人力工时'),
            'spec': '',
            'category': t('labor_hours', lang),
            'unit_price': f'¥{float(lh.get("unit_price", 0)):.2f}',
            'quantity': str(lh.get('hours', 0)),
            'subtotal': f'¥{float(lh.get("total", 0)):.2f}',
        })
    fees_total += labor_total_sum
    if fee_items:
        module_groups.append({
            'name': t('other_fees', lang),
            'items': fee_items,
            'total': fees_total
        })
    
    if module_groups:
        combined_table = doc.add_table(rows=1, cols=8)
        combined_table.style = 'Table Grid'
        headers = ['模块', '项目', '规格', '分类/品牌', '单价', '数量', '小计', '合计']
        for j, header in enumerate(headers):
            combined_table.rows[0].cells[j].text = header
        
        for group in module_groups:
            if not group['items']:
                row = combined_table.add_row()
                row.cells[0].text = group['name']
                row.cells[1].text = '（无物料）'
                continue
            
            first_row = combined_table.add_row()
            first_row.cells[0].text = group['name']
            first_row.cells[1].text = group['items'][0]['item']
            first_row.cells[2].text = group['items'][0]['spec']
            first_row.cells[3].text = group['items'][0]['category']
            first_row.cells[4].text = group['items'][0]['unit_price']
            first_row.cells[5].text = group['items'][0]['quantity']
            first_row.cells[6].text = group['items'][0]['subtotal']
            first_row.cells[7].text = f'¥{group["total"]:.2f}'
            
            if len(group['items']) > 1:
                module_cell = first_row.cells[0]
                for i in range(1, len(group['items'])):
                    next_row = combined_table.add_row()
                    next_row.cells[1].text = group['items'][i]['item']
                    next_row.cells[2].text = group['items'][i]['spec']
                    next_row.cells[3].text = group['items'][i]['category']
                    next_row.cells[4].text = group['items'][i]['unit_price']
                    next_row.cells[5].text = group['items'][i]['quantity']
                    next_row.cells[6].text = group['items'][i]['subtotal']
                    module_cell = module_cell.merge(next_row.cells[0])
                
                total_cell = first_row.cells[7]
                for i in range(1, len(group['items'])):
                    next_row_idx = len(combined_table.rows) - 1
                    total_cell = total_cell.merge(combined_table.rows[next_row_idx].cells[7])
    else:
        doc.add_paragraph('无物料及费用数据')
    
    doc.add_paragraph()
    
    # 汇总信息
    doc.add_heading(t('quote_summary', lang), level=1)
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Table Grid'
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary_data = [
        (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
        (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
        (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
        (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
        (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
        (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
        (f'{t("grand_total", lang)}({currency})', f'{currency_symbol}{grand_total_converted:.2f}')
    ]
    for i, (key, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = value
    
    doc.save(word_path)


def _generate_version_pdf(pdf_path, quotation_id, version_no, data, totals, quotation=None, lang='zh'):
    """生成版本 PDF 文件（复用报价单PDF逻辑）"""
    # 获取币种和汇率
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)
    
    # 获取费用系数（优先用快照中的报价单系数，否则用系统费率）
    fee_rates = data.get('coefficients') or totals.get('fee_rates', {})
    rate_large = fee_rates.get('large', 1.0)
    rate_standard = fee_rates.get('standard', 1.0)
    rate_other = fee_rates.get('other', 1.0)

    # 收集分组数据
    module_groups = []
    for module in data.get('modules', []):
        items = []
        module_total = 0
        for mm in module.get('materials', []):
            mat = Material.query.get(mm.get('material_id'))
            if mat:
                cat = mat.category or 'standard'
                rate = rate_large if cat == 'large' else (rate_standard if cat == 'standard' else rate_other)
                unit_price_with_rate = float(mat.unit_price or 0) * rate
                item_total = unit_price_with_rate * mm.get('quantity', 1)
                module_total += item_total
                items.append({
                    'item': mat.name or '',
                    'spec': mat.spec or '',
                    'category': mat.brand or '',
                    'unit_price': f'¥{unit_price_with_rate:.2f}',
                    'quantity': str(mm.get('quantity', 1)),
                    'subtotal': f'¥{item_total:.2f}',
                })
        if items:
            module_obj = Module.query.get(module.get('id'))
            module_name = get_name(module_obj, lang) if module_obj else module.get('name', t('unnamed_module', lang))
            module_groups.append({
                'name': module_name,
                'items': items,
                'total': module_total
            })
    
    fees_total = sum(float(f.get('amount', 0)) for f in data.get('fees', []))
    fee_items = []
    for fee in data.get('fees', []):
        location_text = '厂内' if fee.get('location') == 'internal' else ('厂外' if fee.get('location') == 'external' else '')
        fee_items.append({
            'item': get_fee_type_name(fee.get('fee_type', ''), lang) or fee.get('name', ''),
            'spec': location_text,
            'category': get_fee_type_name(fee.get('fee_type', ''), lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f"¥{float(fee.get('amount', 0)):.2f}",
        })
    # 人力工时明细加入"其他费用"分组
    labor_total_sum = 0.0
    for lh in data.get('labor_hours', []):
        labor_total_sum += float(lh.get('total', 0))
        fee_items.append({
            'item': lh.get('name', t('labor_hours', lang)),
            'spec': '',
            'category': t('labor_hours', lang),
            'unit_price': f'¥{float(lh.get("unit_price", 0)):.2f}',
            'quantity': str(lh.get('hours', 0)),
            'subtotal': f'¥{float(lh.get("total", 0)):.2f}',
        })
    fees_total += labor_total_sum
    if fee_items:
        module_groups.append({
            'name': t('other_fees', lang),
            'items': fee_items,
            'total': fees_total
        })
    
    buffer = BytesIO()
    
    class PDF(FPDF):
        def header(self):
            pass
        def footer(self):
            pass
    
    pdf = PDF()
    pdf.add_font('SimHei', '', FONT_REGULAR, uni=True)
    pdf.add_font('SimHei', 'B', FONT_BOLD, uni=True)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # 标题
    pdf.set_font('SimHei', 'B', 16)
    pdf.cell(0, 12, data.get('name', '报价单'), 0, 1, 'C')
    pdf.ln(5)
    
    # 基本信息
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('basic_info', lang), 0, 1)
    pdf.set_font('SimHei', '', 10)
    
    business_owner_name = ''
    if data.get('business_owner_id'):
        owner = User.query.get(data.get('business_owner_id'))
        if owner:
            business_owner_name = owner.real_name
    basic_info = [
        (t('project_name', lang), data.get('name', '')),
        (t('scheme_no', lang), data.get('scheme_no', '')),
        (t('business_owner', lang), business_owner_name),
        (t('version_no', lang), f'v{version_no}'),
        (t('currency', lang), currency)
    ]
    for key, value in basic_info:
        pdf.set_font('SimHei', 'B', 9)
        pdf.cell(40, 6, key, 1)
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 6, str(value), 1, 1)
    
    pdf.ln(5)
    
    # 物料及费用清单 (PDF)
    if module_groups:
        col_widths = [22, 38, 26, 20, 22, 12, 26, 24]  # 总约190
        headers = [t('module_col', lang), t('item_col', lang), t('spec_col', lang),
                   t('category_col', lang), t('unit_price_col', lang), t('qty_col', lang),
                   t('subtotal_col', lang), t('total_col', lang)]
        
        def draw_table_header():
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            for i, h in enumerate(headers):
                pdf.cell(col_widths[i], 6, h, 1, 0, 'C', True)
            pdf.ln()
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
        
        def draw_group(y_start, group):
            item_count = len(group['items'])
            row_h = 5
            y_bottom = y_start + item_count * row_h
            
            # 使用PDF实际的左边距（默认10）
            x_start = pdf.l_margin if pdf.l_margin else 10
            x_module = x_start
            x_data = x_start + col_widths[0]
            x_total = x_start + col_widths[0] + sum(col_widths[1:7])
            x_end = x_start + sum(col_widths)
            
            # 模块列外边框
            pdf.rect(x_module, y_start, col_widths[0], item_count * row_h)
            # 合计列外边框
            pdf.rect(x_total, y_start, col_widths[7], item_count * row_h)
            
            # 数据区域边框
            pdf.line(x_data, y_start, x_data, y_bottom)
            pdf.line(x_total, y_start, x_total, y_bottom)
            
            # 列竖向分割线
            x_cur = x_data
            for j in range(1, 7):
                x_line = x_cur + col_widths[j]
                pdf.line(x_line, y_start, x_line, y_bottom)
                x_cur += col_widths[j]
            
            # 行水平分割线
            for i in range(1, item_count):
                y_line = y_start + i * row_h
                pdf.line(x_data, y_line, x_total, y_line)
            
            # 顶边和底边
            pdf.line(x_data, y_start, x_total, y_start)
            pdf.line(x_data, y_bottom, x_total, y_bottom)
            
            # 文字 - 模块列使用 multi_cell 自动换行
            text = group['name']
            line_height = 4
            lines_wrapped = pdf.multi_cell(col_widths[0], line_height, text, border=0, align='C', split_only=True)
            text_height = len(lines_wrapped) * line_height
            cell_height = item_count * row_h
            y_text = y_start + (cell_height - text_height) / 2
            pdf.set_xy(x_module, y_text)
            pdf.multi_cell(col_widths[0], line_height, text, 0, 'C')
            # 合计列
            pdf.set_xy(x_total, y_start + (item_count * row_h - 5) / 2)
            pdf.cell(col_widths[7], 5, f'¥{group["total"]:.2f}', 0, 0, 'C')
            
            # 数据单元格内容
            for i, item in enumerate(group['items']):
                y_row = y_start + i * row_h
                x_cur = x_data
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[1], row_h, str(item['item'])[:17], 0)
                x_cur += col_widths[1]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[2], row_h, str(item['spec'])[:12], 0)
                x_cur += col_widths[2]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[3], row_h, str(item['category'])[:9], 0)
                x_cur += col_widths[3]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[4], row_h, str(item['unit_price']), 0, 0, 'R')
                x_cur += col_widths[4]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[5], row_h, str(item['quantity']), 0, 0, 'C')
                x_cur += col_widths[5]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[6], row_h, str(item['subtotal']), 0, 0, 'R')
            
            # 移动到下一行
            pdf.set_y(y_bottom)
            return y_bottom
        
        def check_page_break(y_current, item_count, row_h=5):
            """检查是否需要换页，留出足够空间给当前模块"""
            try:
                y_current = float(y_current) if y_current is not None else 0
            except (ValueError, TypeError):
                y_current = 0
            pdf_h = float(pdf.h) if pdf.h else 297.0
            pdf_bm = float(pdf.b_margin) if pdf.b_margin else 10.0
            page_height = pdf_h - pdf_bm
            needed_height = item_count * row_h + 10  # 10为表头高度
            if y_current + needed_height > page_height:
                pdf.add_page()
                draw_table_header()
                new_y = pdf.get_y()
                return float(new_y) if new_y is not None else (pdf_h - pdf_bm - 10)
            return y_current
        
        # 开始绘制表格
        pdf.ln(5)
        pdf.set_font('SimHei', 'B', 11)
        pdf.cell(0, 8, t('material_list', lang), 0, 1)
        draw_table_header()
        
        y_start = pdf.get_y()
        for group in module_groups:
            if not group['items']:
                pdf.cell(sum(col_widths), 5, f'{group["name"]} - （无物料）', 1, 1, 'L')
                y_start = pdf.get_y()
                continue
            
            item_count = len(group['items'])
            # 检查是否需要换页
            y_start = check_page_break(y_start, item_count)
            
            y_end = draw_group(y_start, group)
            if y_end is None:
                y_end = y_start + len(group['items']) * 5
            y_start = y_end
        
        # 移动到下一页
        if y_start + 10 > (pdf.h if pdf.h else 297) - (pdf.b_margin if pdf.b_margin else 10):
            pdf.add_page()
    else:
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 5, '无物料及费用数据', 0, 1)
    
    pdf.ln(5)
    
    # 报价汇总
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('quote_summary', lang), 0, 1)
    
    pdf.set_font('SimHei', 'B', 9)
    pdf.set_fill_color(68, 114, 196)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(100, 6, t('project_col', lang), 1, 0, 'C', True)
    pdf.cell(0, 6, t('amount_col', lang), 1, 1, 'C', True)
    
    pdf.set_text_color(0, 0, 0)
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary_items = [
        (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
        (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
        (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
        (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
        (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
        (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
    ]
    pdf.set_font('SimHei', '', 9)
    for label, value in summary_items:
        pdf.cell(100, 5, label, 1)
        pdf.cell(0, 5, value, 1, 1, 'R')
    
    # 最终报价（高亮）
    pdf.set_font('SimHei', 'B', 10)
    pdf.set_fill_color(226, 239, 218)
    pdf.cell(100, 7, f"{t('grand_total', lang)}({currency})", 1, 0, 'C', True)
    pdf.cell(0, 7, f'{currency_symbol}{grand_total_converted:.2f}', 1, 1, 'R', True)
    
    # 输出
    pdf_bytes = pdf.output()
    with open(pdf_path, 'wb') as f:
        f.write(pdf_bytes)

@export_bp.route('/quotations/<int:quotation_id>/export/word', methods=['GET'])
def export_word(quotation_id):
    """导出 Word 格式"""
    data = get_quotation_with_details(quotation_id)
    if not data:
        return jsonify({'error': 'Quotation not found'}), 404
    
    quotation, modules, fees, labor_hours = data
    totals = calculate_totals_with_rates(quotation, modules, fees, labor_hours)
    
    # 使用报价单的币种
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    
    # 获取汇率
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    
    # 转换最终报价
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)
    
    doc = Document()
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('报价单', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('基本信息', level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        (t('project_name', lang), quotation.name or ''),
        (t('scheme_no', lang), quotation.scheme_no or ''),
        (t('business_owner', lang), quotation.business_owner.real_name if quotation.business_owner else ''),
        ('创建时间', quotation.created_at.strftime('%Y-%m-%d %H:%M') if quotation.created_at else ''),
        (t('currency', lang), currency)
    ]
    for i, (key, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # 整合表格：模块及物料 + 其他费用
    doc.add_heading('物料及费用清单', level=1)
    
    # 收集分组数据
    module_groups = []
    
    # 按模块分组
    for module in modules:
        items = []
        module_total = 0
        if module.materials:
            for mm in module.materials:
                mat = mm.material
                if mat:
                    item_total = float(mat.unit_price or 0) * mm.quantity
                    module_total += item_total
                    items.append({
                        'item': mat.name or '',
                        'spec': mat.spec or '',
                        'category': mat.brand or '',
                        'unit_price': f'¥{float(mat.unit_price or 0):.2f}',
                        'quantity': str(mm.quantity),
                        'subtotal': f'¥{item_total:.2f}',
                    })
        module_groups.append({
            'name': get_name(module, lang),
            'items': items,
            'total': module_total
        })
    
    # 费用 + 人力工时单独一组
    fees_total = sum(float(f.amount or 0) for f in fees)
    labor_total = sum(float(l.total or 0) for l in labor_hours)
    fee_items = []
    for fee in fees:
        location_text = '厂内' if fee.location == 'internal' else ('厂外' if fee.location == 'external' else fee.location or '')
        fee_items.append({
            'item': get_fee_type_name(fee.fee_type, lang) or '',
            'spec': location_text,
            'category': get_fee_type_name(fee.fee_type, lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f'¥{float(fee.amount or 0):.2f}',
        })
    for l in labor_hours:
        fee_items.append({
            'item': l.name or '人力工时',
            'spec': f'{l.hours}h',
            'category': '人力工时',
            'unit_price': f'¥{float(l.unit_price or 0):.2f}/h',
            'quantity': '1',
            'subtotal': f'¥{float(l.total or 0):.2f}',
        })
    if fee_items:
        module_groups.append({
            'name': '其他费用',
            'items': fee_items,
            'total': fees_total + labor_total
        })
    
    if module_groups:
        # 创建整合表格：模块|项目|规格|分类/品牌|单价|数量|小计|合计
        combined_table = doc.add_table(rows=1, cols=8)
        combined_table.style = 'Table Grid'
        headers = ['模块', '项目', '规格', '分类/品牌', '单价', '数量', '小计', '合计']
        for j, header in enumerate(headers):
            combined_table.rows[0].cells[j].text = header
        
        for group in module_groups:
            if not group['items']:
                row = combined_table.add_row()
                row.cells[0].text = group['name']
                row.cells[1].text = '（无物料）'
                continue
            
            # 第一行用于合并模块和合计
            first_row = combined_table.add_row()
            first_row.cells[0].text = group['name']
            first_row.cells[1].text = group['items'][0]['item']
            first_row.cells[2].text = group['items'][0]['spec']
            first_row.cells[3].text = group['items'][0]['category']
            first_row.cells[4].text = group['items'][0]['unit_price']
            first_row.cells[5].text = group['items'][0]['quantity']
            first_row.cells[6].text = group['items'][0]['subtotal']
            first_row.cells[7].text = f'¥{group["total"]:.2f}'
            
            # 合并模块列 (垂直合并)
            if len(group['items']) > 1:
                # 合并模块列
                module_cell = first_row.cells[0]
                for i in range(1, len(group['items'])):
                    next_row = combined_table.add_row()
                    next_row.cells[1].text = group['items'][i]['item']
                    next_row.cells[2].text = group['items'][i]['spec']
                    next_row.cells[3].text = group['items'][i]['category']
                    next_row.cells[4].text = group['items'][i]['unit_price']
                    next_row.cells[5].text = group['items'][i]['quantity']
                    next_row.cells[6].text = group['items'][i]['subtotal']
                    # 合并模块列
                    module_cell = module_cell.merge(next_row.cells[0])
                
                # 合并合计列
                total_cell = first_row.cells[7]
                for i in range(1, len(group['items'])):
                    next_row_idx = len(combined_table.rows) - 1
                    total_cell = total_cell.merge(combined_table.rows[next_row_idx].cells[7])
    else:
        doc.add_paragraph('无物料及费用数据')
    
    doc.add_paragraph()
    
    # 汇总信息
    doc.add_heading(t('quote_summary', lang), level=1)
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Table Grid'
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary_data = [
        (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
        (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
        (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
        (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
        (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
        (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
        (f'{t("grand_total", lang)}({currency})', f'{currency_symbol}{grand_total_converted:.2f}')
    ]
    for i, (key, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = value
    
    # 保存文件
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f'quotation_{quotation_id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@export_bp.route('/quotations/<int:quotation_id>/export/excel', methods=['GET'])
def export_excel(quotation_id):
    """导出 Excel 格式"""
    data = get_quotation_with_details(quotation_id)
    if not data:
        return jsonify({'error': 'Quotation not found'}), 404
    
    quotation, modules, fees, labor_hours = data
    totals = calculate_totals_with_rates(quotation, modules, fees, labor_hours)
    
    # 使用报价单的币种
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    
    # 获取汇率
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    
    # 转换最终报价
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)
    
    wb = Workbook()
    wb.remove(wb.active)  # 删除默认 Sheet
    
    # 样式定义
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Sheet 1: 基本信息
    ws1 = wb.create_sheet('基本信息')
    
    ws1.append(['报价单基本信息'])
    ws1.merge_cells('A1:D1')
    ws1['A1'].font = Font(bold=True, size=14)
    ws1['A1'].alignment = Alignment(horizontal='center')
    
    ws1.append([])
    basic_info = [
        (t('project_name', lang), quotation.name or ''),
        (t('scheme_no', lang), quotation.scheme_no or ''),
        (t('business_owner', lang), quotation.business_owner.real_name if quotation.business_owner else ''),
        ('创建时间', quotation.created_at.strftime('%Y-%m-%d %H:%M') if quotation.created_at else ''),
        (t('currency', lang), currency)
    ]
    for key, value in basic_info:
        ws1.append([key, value])
    
    ws1.append([])
    ws1.append(['费用系数'])
    ws1.merge_cells(f'A{ws1.max_row}:B{ws1.max_row}')
    ws1[f'A{ws1.max_row}'].font = Font(bold=True)
    
    rate_info = [
        ('大件系数', f"{totals['fee_rates'].get('large', 1.0)}x"),
        ('普通件系数', f"{totals['fee_rates'].get('standard', 1.0)}x"),
        ('其他件系数', f"{totals['fee_rates'].get('other', 1.0)}x")
    ]
    for key, value in rate_info:
        ws1.append([key, value])
    
    ws1.append([])
    ws1.append(['报价汇总'])
    ws1.merge_cells(f'A{ws1.max_row}:B{ws1.max_row}')
    ws1[f'A{ws1.max_row}'].font = Font(bold=True)
    
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary = [
        ('物料合计', f'¥{totals["material_total_with_rates"]:.2f}'),
        ('费用合计', f'¥{fees_with_labor:.2f}'),
        ('利润率', f'{totals["profit_rate"] * 100:.1f}%'),
        ('含利润率小计', f'¥{totals["subtotal_with_profit"]:.2f}'),
        ('税率', f'{totals["tax_rate"] * 100:.0f}%'),
        ('税额', f'¥{totals["tax_amount"]:.2f}'),
        (f'最终报价({currency})', f'{currency_symbol}{grand_total_converted:.2f}')
    ]
    for key, value in summary:
        ws1.append([key, value])
    
    # 调整列宽
    ws1.column_dimensions['A'].width = 18
    ws1.column_dimensions['B'].width = 30
    
    # Sheet 2: 整合物料及费用清单
    ws2 = wb.create_sheet('物料及费用清单')
    
    # 添加表头 - 8列
    headers = ['模块', '项目', '规格', '分类/品牌', '单价', '数量', '小计', '合计']
    ws2.append(headers)
    for cell in ws2[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border
    
    row_num = 2
    merge_info = []  # 记录需要合并的信息：(start_row, end_row, col_letter)
    
    for group in module_groups:
        if not group['items']:
            ws2.append([group['name'], '（无物料）', '', '', '', '', '', ''])
            for cell in ws2[row_num]:
                cell.border = thin_border
            row_num += 1
            continue
        
        start_row = row_num
        for i, item in enumerate(group['items']):
            ws2.append([
                group['name'] if i == 0 else '',
                item['item'],
                item['spec'],
                item['category'],
                float(item['unit_price'].replace('¥', '')) if item['unit_price'] else 0,
                int(item['quantity']),
                float(item['subtotal'].replace('¥', '')) if item['subtotal'] else 0,
                group['total'] if i == 0 else ''
            ])
            for cell in ws2[row_num]:
                cell.border = thin_border
            row_num += 1
        
        # 记录合并信息
        if len(group['items']) > 1:
            merge_info.append((start_row, row_num - 1, 'A'))  # 模块列
            merge_info.append((start_row, row_num - 1, 'H'))  # 合计列
    
    # 执行合并
    for start, end, col in merge_info:
        if start < end:
            ws2.merge_cells(f'{col}{start}:{col}{end}')
            # 居中对齐
            cell = ws2[f'{col}{start}']
            cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # 调整列宽
    col_widths = [16, 24, 18, 14, 12, 8, 12, 16]
    for i, width in enumerate(col_widths):
        ws2.column_dimensions[get_column_letter(i + 1)].width = width
    
    # 保存文件
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    filename = f'quotation_{quotation_id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.xlsx'
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )


# ==================== 版本导出（指定版本号） ====================

def _build_export_data_from_snapshot(snapshot_data_str):
    """从 snapshot_data 构建脱敏导出数据（用于高效渲染/导出）"""
    import json
    data = json.loads(snapshot_data_str) if isinstance(snapshot_data_str, str) else snapshot_data_str
    # 已在 normalize 后结构化为 {quotation, modules, fees, labor_hours}
    return data


@export_bp.route('/quotations/<int:quotation_id>/versions/<int:version_no>/export/excel', methods=['GET'])
def export_version_excel(quotation_id, version_no):
    """导出版本 Excel（指定版本号）"""
    version = VersionSnapshot.query.filter_by(quotation_id=quotation_id, version_no=version_no).first()
    if not version:
        return jsonify({'error': 'Version not found'}), 404

    # 如果有预生成文件，直接返回（暂未实现，动态生成）
    data = get_version_snapshot_data(quotation_id, version_no)
    if not data:
        return jsonify({'error': 'Version data not found'}), 404

    quotation = Quotation.query.get(quotation_id)
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total = convert_currency(
        calculate_version_totals(data)['grand_total'], currency, exchange_rates)

    wb = Workbook()
    wb.remove(wb.active)

    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                         top=Side(style='thin'), bottom=Side(style='thin'))

    # Sheet 1: 基本信息
    ws1 = wb.create_sheet('基本信息')
    ws1.append(['报价单基本信息'])
    ws1.merge_cells('A1:D1')
    ws1['A1'].font = Font(bold=True, size=14)
    ws1['A1'].alignment = Alignment(horizontal='center')
    ws1.append([])
    info = [
        ('项目名称', data.get('name', '')),
        ('方案编号', data.get('scheme_no', '')),
        ('版本号', f'v{version_no}'),
        ('币种', currency),
    ]
    for k, v in info:
        ws1.append([k, v])
    ws1.column_dimensions['A'].width = 18
    ws1.column_dimensions['B'].width = 30

    # Sheet 2: 物料及费用清单
    ws2 = wb.create_sheet('物料及费用清单')
    headers = ['模块', '项目', '规格', '分类/品牌', '单价', '数量', '小计', '合计']
    ws2.append(headers)
    for cell in ws2[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    modules = data.get('modules', [])
    for mod in modules:
        mod_name = mod.get('name', '未命名模块')
        materials = mod.get('materials', [])
        if not materials:
            ws2.append([mod_name, '（无物料）', '', '', '', '', '', ''])
            continue
        for i, m in enumerate(materials):
            ws2.append([
                mod_name if i == 0 else '',
                m.get('name') or m.get('item', ''),
                m.get('spec', ''),
                m.get('brand', ''),
                m.get('unit_price', 0),
                m.get('quantity', 0),
                m.get('subtotal', 0) if isinstance(m.get('subtotal'), (int, float)) else 0,
                '',
            ])
        # 模块合计行（取第一条物料的小计之和）
        mod_total = sum(
            float(m.get('subtotal', 0)) if isinstance(m.get('subtotal'), (int, float)) else 0
            for m in materials
        )
        last_row = ws2.max_row
        ws2[f'H{last_row}'] = mod_total
        ws2.merge_cells(f'A{last_row}:G{last_row}')
        ws2[f'A{last_row}'].alignment = Alignment(horizontal='right')

    # 调整列宽
    for i, w in enumerate([15, 20, 20, 15, 12, 8, 12, 12]):
        ws2.column_dimensions[get_column_letter(i + 1)].width = w

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    filename = f'{data.get("name", "报价单")}_v{version_no}.xlsx'
    return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    as_attachment=True, download_name=filename)


@export_bp.route('/quotations/<int:quotation_id>/export/pdf', methods=['GET'])
def export_pdf(quotation_id):
    """导出 PDF 格式，支持中英文 ?lang=en """
    lang = request.args.get('lang', 'zh')
    data = get_quotation_with_details(quotation_id)
    if not data:
        return jsonify({'error': 'Quotation not found'}), 404
    
    quotation, modules, fees, labor_hours = data
    totals = calculate_totals_with_rates(quotation, modules, fees, labor_hours)
    
    # 使用报价单的币种
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    
    # 获取汇率
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    
    # 转换最终报价
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)
    
    # 获取系数
    coeff = (quotation.coefficients if hasattr(quotation, 'coefficients') and quotation.coefficients else
             {fr.category: float(fr.rate) for fr in FeeRate.query.all()})
    rate_large = float(coeff.get('large', 1.0))
    rate_standard = float(coeff.get('standard', 1.0))
    rate_other = float(coeff.get('other', 1.0))
    
    doc = Document()
    module_groups = []
    for module in modules:
        items = []
        module_total = 0
        if module.materials:
            for mm in module.materials:
                mat = mm.material
                if mat:
                    cat = mat.category or 'standard'
                    rate = rate_large if cat == 'large' else (rate_standard if cat == 'standard' else rate_other)
                    unit_price_with_rate = float(mat.unit_price or 0) * rate
                    item_total = unit_price_with_rate * mm.quantity
                    module_total += item_total
                    items.append({
                        'item': mat.name or '',
                        'spec': mat.spec or '',
                        'category': mat.brand or '',
                        'unit_price': f'¥{unit_price_with_rate:.2f}',
                        'quantity': str(mm.quantity),
                        'subtotal': f'¥{item_total:.2f}',
                    })
        module_groups.append({
            'name': get_name(module, lang),
            'items': items,
            'total': module_total
        })
    
    fees_total = sum(float(f.amount or 0) for f in fees)
    labor_total = sum(float(l.total or 0) for l in labor_hours)
    fee_items = []
    for fee in fees:
        location_text = '厂内' if fee.location == 'internal' else ('厂外' if fee.location == 'external' else fee.location or '')
        fee_items.append({
            'item': get_fee_type_name(fee.fee_type, lang) or '',
            'spec': location_text,
            'category': get_fee_type_name(fee.fee_type, lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f'¥{float(fee.amount or 0):.2f}',
        })
    for l in labor_hours:
        fee_items.append({
            'item': l.name or t('labor_hours', lang),
            'spec': f"{l.hours}{t('unit_h', lang)}",
            'category': t('labor_hours', lang),
            'unit_price': f'¥{float(l.unit_price or 0):.2f}/h',
            'quantity': '1',
            'subtotal': f'¥{float(l.total or 0):.2f}',
        })
    if fee_items:
        module_groups.append({
            'name': t('other_fees', lang),
            'items': fee_items,
            'total': fees_total + labor_total
        })
    
    class PDF(FPDF):
        def header(self):
            pass
        def footer(self):
            pass
    
    pdf = PDF()
    pdf.add_font('SimHei', '', FONT_REGULAR, uni=True)
    pdf.add_font('SimHei', 'B', FONT_BOLD, uni=True)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # 标题
    pdf.set_font('SimHei', 'B', 16)
    pdf.cell(0, 12, t('quote', lang), 0, 1, 'C')
    pdf.ln(5)
    
    # 基本信息
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('basic_info', lang), 0, 1)
    pdf.set_font('SimHei', '', 10)
    
    basic_info = [
        (t('project_name', lang), quotation.name or ''),
        (t('scheme_no', lang), quotation.scheme_no or ''),
        (t('business_owner', lang), quotation.business_owner.real_name if quotation.business_owner else ''),
        (t('created_at', lang), quotation.created_at.strftime('%Y-%m-%d %H:%M') if quotation.created_at else ''),
        (t('currency', lang), currency)
    ]
    for key, value in basic_info:
        pdf.set_font('SimHei', 'B', 9)
        pdf.cell(40, 6, key, 1)
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 6, str(value), 1, 1)
    
    pdf.ln(5)
    
    # 物料及费用清单 (PDF)
    if module_groups:
        col_widths = [22, 38, 26, 20, 22, 12, 26, 24]  # 总约190
        headers = [t('module_col', lang), t('item_col', lang), t('spec_col', lang),
                   t('category_col', lang), t('unit_price_col', lang), t('qty_col', lang),
                   t('subtotal_col', lang), t('total_col', lang)]
        
        def draw_table_header():
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            for i, h in enumerate(headers):
                pdf.cell(col_widths[i], 6, h, 1, 0, 'C', True)
            pdf.ln()
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
        
        def draw_group(y_start, group):
            item_count = len(group['items'])
            row_h = 5
            y_bottom = y_start + item_count * row_h
            
            # 使用PDF实际的左边距（默认10）
            x_start = pdf.l_margin if pdf.l_margin else 10
            x_module = x_start
            x_data = x_start + col_widths[0]
            x_total = x_start + col_widths[0] + sum(col_widths[1:7])
            x_end = x_start + sum(col_widths)
            
            # 模块列外边框
            pdf.rect(x_module, y_start, col_widths[0], item_count * row_h)
            # 合计列外边框
            pdf.rect(x_total, y_start, col_widths[7], item_count * row_h)
            
            # 数据区域边框
            pdf.line(x_data, y_start, x_data, y_bottom)
            pdf.line(x_total, y_start, x_total, y_bottom)
            
            # 列竖向分割线
            x_cur = x_data
            for j in range(1, 7):
                x_line = x_cur + col_widths[j]
                pdf.line(x_line, y_start, x_line, y_bottom)
                x_cur += col_widths[j]
            
            # 行水平分割线
            for i in range(1, item_count):
                y_line = y_start + i * row_h
                pdf.line(x_data, y_line, x_total, y_line)
            
            # 顶边和底边
            pdf.line(x_data, y_start, x_total, y_start)
            pdf.line(x_data, y_bottom, x_total, y_bottom)
            
            # 文字 - 模块列使用 multi_cell 自动换行
            text = group['name']
            line_height = 4
            lines_wrapped = pdf.multi_cell(col_widths[0], line_height, text, border=0, align='C', split_only=True)
            text_height = len(lines_wrapped) * line_height
            cell_height = item_count * row_h
            y_text = y_start + (cell_height - text_height) / 2
            pdf.set_xy(x_module, y_text)
            pdf.multi_cell(col_widths[0], line_height, text, 0, 'C')
            # 合计列
            pdf.set_xy(x_total, y_start + (item_count * row_h - 5) / 2)
            pdf.cell(col_widths[7], 5, f'¥{group["total"]:.2f}', 0, 0, 'C')
            
            # 数据单元格内容
            for i, item in enumerate(group['items']):
                y_row = y_start + i * row_h
                x_cur = x_data
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[1], row_h, str(item['item'])[:17], 0)
                x_cur += col_widths[1]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[2], row_h, str(item['spec'])[:12], 0)
                x_cur += col_widths[2]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[3], row_h, str(item['category'])[:9], 0)
                x_cur += col_widths[3]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[4], row_h, str(item['unit_price']), 0, 0, 'R')
                x_cur += col_widths[4]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[5], row_h, str(item['quantity']), 0, 0, 'C')
                x_cur += col_widths[5]
                pdf.set_xy(x_cur, y_row)
                pdf.cell(col_widths[6], row_h, str(item['subtotal']), 0, 0, 'R')
            
            # 移动到下一行
            pdf.set_y(y_bottom)
            return y_bottom
        
        def check_page_break(y_current, item_count, row_h=5):
            """检查是否需要换页，留出足够空间给当前模块"""
            try:
                y_current = float(y_current) if y_current is not None else 0
            except (ValueError, TypeError):
                y_current = 0
            pdf_h = float(pdf.h) if pdf.h else 297.0
            pdf_bm = float(pdf.b_margin) if pdf.b_margin else 10.0
            page_height = pdf_h - pdf_bm
            needed_height = item_count * row_h + 10  # 10为表头高度
            if y_current + needed_height > page_height:
                pdf.add_page()
                draw_table_header()
                new_y = pdf.get_y()
                return float(new_y) if new_y is not None else (pdf_h - pdf_bm - 10)
            return y_current
        
        # 开始绘制表格
        pdf.ln(5)
        pdf.set_font('SimHei', 'B', 11)
        pdf.cell(0, 8, t('material_list', lang), 0, 1)
        draw_table_header()
        
        y_start = pdf.get_y()
        for group in module_groups:
            if not group['items']:
                pdf.cell(sum(col_widths), 5, f'{group["name"]} - {t("no_materials", lang)}', 1, 1, 'L')
                y_start = pdf.get_y()
                continue
            
            item_count = len(group['items'])
            # 检查是否需要换页
            y_start = check_page_break(y_start, item_count)
            
            y_end = draw_group(y_start, group)
            if y_end is None:
                y_end = y_start + len(group['items']) * 5
            y_start = y_end
        
        # 移动到下一页
        if y_start + 10 > (pdf.h if pdf.h else 297) - (pdf.b_margin if pdf.b_margin else 10):
            pdf.add_page()
    else:
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 5, '无物料及费用数据', 0, 1)
    
    pdf.ln(5)
    
    # 报价汇总
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('quote_summary', lang), 0, 1)
    
    pdf.set_font('SimHei', 'B', 9)
    pdf.set_fill_color(68, 114, 196)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(100, 6, t('project_col', lang), 1, 0, 'C', True)
    pdf.cell(0, 6, t('amount_col', lang), 1, 1, 'C', True)
    
    pdf.set_text_color(0, 0, 0)
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary_items = [
        (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
        (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
        (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
        (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
        (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
        (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
    ]
    pdf.set_font('SimHei', '', 9)
    for label, value in summary_items:
        pdf.cell(100, 5, label, 1)
        pdf.cell(0, 5, value, 1, 1, 'R')
    
    # 最终报价（高亮）
    pdf.set_font('SimHei', 'B', 10)
    pdf.set_fill_color(226, 239, 218)
    pdf.cell(100, 7, f"{t('grand_total', lang)}({currency})", 1, 0, 'C', True)
    pdf.cell(0, 7, f'{currency_symbol}{grand_total_converted:.2f}', 1, 1, 'R', True)
    
    # 输出到临时文件
    pdf_bytes = pdf.output()
    buffer = BytesIO(pdf_bytes)
    filename = f'quotation_{quotation_id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.pdf'
    return send_file(buffer, mimetype='application/pdf', as_attachment=True, download_name=filename)


# ==================== 版本导出 ====================

def get_version_snapshot_data(quotation_id, version_no):
    """获取版本快照数据并重建完整结构"""
    version = VersionSnapshot.query.filter_by(
        quotation_id=quotation_id,
        version_no=version_no
    ).first()
    if not version:
        return None
    
    import json
    data = json.loads(version.snapshot_data)
    
    # 重建物料信息（从快照的material_id查询完整物料数据）
    for module in data.get('modules', []):
        materials = module.get('materials', [])
        for mm in materials:
            mat = Material.query.get(mm.get('material_id'))
            if mat:
                mm['name'] = mat.name or ''
                mm['brand'] = mat.brand or ''
                mm['spec'] = mat.spec or ''
                mm['unit_price'] = mat.unit_price or 0
                mm['category'] = mat.category or 'standard'
    
    # 重建费用信息
    for fee in data.get('fees', []):
        fee['name'] = fee.get('fee_type', '') or fee.get('name', '')
        fee['amount'] = float(fee.get('amount', 0))
        fee['location'] = fee.get('position', 'internal')
    
    return data


def calculate_version_totals(data):
    """计算版本数据的汇总"""
    modules = data.get('modules', [])
    fees = data.get('fees', [])
    
    # 分类统计物料
    category_totals = {'large': 0, 'standard': 0, 'other': 0}
    for module in modules:
        for mm in module.get('materials', []):
            cat = mm.get('category', 'standard')
            amount = float(mm.get('unit_price', 0) or 0) * float(mm.get('quantity', 0) or 0)
            category_totals[cat] = category_totals.get(cat, 0) + amount
    
    # 获取费用系数（优先用快照中的报价单系数，否则用系统费率）
    coeff = data.get('coefficients')
    if coeff:
        rate_large = coeff.get('large', 1.0)
        rate_standard = coeff.get('standard', 1.0)
        rate_other = coeff.get('other', 1.0)
        fee_rates = coeff
    else:
        fee_rates = {fr.category: float(fr.rate) for fr in FeeRate.query.all()}
        rate_large = fee_rates.get('large', 1.0)
        rate_standard = fee_rates.get('standard', 1.0)
        rate_other = fee_rates.get('other', 1.0)
    
    material_with_rates = (
        category_totals.get('large', 0) * rate_large +
        category_totals.get('standard', 0) * rate_standard +
        category_totals.get('other', 0) * rate_other
    )

    material_total = sum(category_totals.values())
    fees_total = sum(float(fee.get('amount', 0)) for fee in fees)
    labor_total = sum(float(l.get('total', 0)) for l in data.get('labor_hours', []))
    subtotal = material_with_rates + fees_total + labor_total

    # 获取对外利润率（从报价单或根级获取）
    profit_rate = 0.0
    if data.get('profit_rate') is not None:
        profit_rate = float(data.get('profit_rate', 0))
    elif data.get('quotation', {}).get('profit_rate') is not None:
        profit_rate = float(data['quotation'].get('profit_rate', 0))
    profit_amount = subtotal * profit_rate
    subtotal_with_profit = subtotal + profit_amount

    # 获取税率（从报价单或根级获取）
    tax_rate = 0.13
    if data.get('tax_rate') is not None:
        tax_rate = float(data.get('tax_rate', 0.13))
    elif data.get('quotation', {}).get('tax_rate') is not None:
        tax_rate = float(data['quotation'].get('tax_rate', 0.13))
    grand_total = subtotal_with_profit * (1 + tax_rate)
    tax_amount = grand_total - subtotal_with_profit
    
    return {
        'material_total': material_total,
        'material_total_with_rates': material_with_rates,
        'fees_total': fees_total,
        'labor_total': labor_total,
        'subtotal': subtotal,
        'profit_rate': profit_rate,
        'profit_amount': profit_amount,
        'subtotal_with_profit': subtotal_with_profit,
        'tax_rate': tax_rate,
        'tax_amount': tax_amount,
        'grand_total': grand_total,
        'fee_rates': fee_rates
    }


@export_bp.route('/quotations/<int:quotation_id>/versions/<int:version_no>/export/word', methods=['GET'])
def export_version_word(quotation_id, version_no):
    """导出版本 Word 格式（直接下载预生成文件）"""
    version = VersionSnapshot.query.filter_by(
        quotation_id=quotation_id,
        version_no=version_no
    ).first()
    if not version:
        return jsonify({'error': 'Version not found'}), 404
    
    # 如果有预生成的文件，直接返回
    if version.word_file:
        import os
        if os.path.exists(version.word_file):
            # snapshot_data 是 JSON 字符串
            import json
            name = '报价单'
            if version.snapshot_data:
                try:
                    data = json.loads(version.snapshot_data) if isinstance(version.snapshot_data, str) else version.snapshot_data
                    name = data.get('name', '报价单')
                except:
                    pass
            filename = f'{name}_v{version_no}.docx'
            return send_file(
                version.word_file,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=filename
            )
    
    # 如果没有预生成文件，动态生成
    data = get_version_snapshot_data(quotation_id, version_no)
    if not data:
        return jsonify({'error': 'Version data not found'}), 404
    
    # 获取报价单信息（包含币种）
    quotation = Quotation.query.get(quotation_id)
    
    totals = calculate_version_totals(data)
    # 使用报价单的币种
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)
    
    doc = Document()
    title = doc.add_heading(f'{data.get("name", "报价单")}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('基本信息', level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    business_owner_name = ''
    if data.get('business_owner_id'):
        owner = User.query.get(data.get('business_owner_id'))
        if owner:
            business_owner_name = owner.real_name
    info_data = [
        (t('project_name', lang), data.get('name', '')),
        (t('scheme_no', lang), data.get('scheme_no', '')),
        (t('business_owner', lang), business_owner_name),
        (t('version_no', lang), f'v{version_no}'),
        (t('currency', lang), currency)
    ]
    for i, (key, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # 整合表格：模块及物料 + 其他费用
    doc.add_heading('物料及费用清单', level=1)
    
    module_groups = []
    
    # 按模块分组
    for module in data.get('modules', []):
        items = []
        module_total = 0
        for mm in module.get('materials', []):
            item_total = float(mm.get('unit_price', 0) or 0) * float(mm.get('quantity', 0) or 0)
            module_total += item_total
            items.append({
                'item': mm.get('name', ''),
                'spec': mm.get('spec', ''),
                'category': mm.get('brand', ''),
                'unit_price': f'¥{float(mm.get("unit_price", 0) or 0):.2f}',
                'quantity': str(mm.get('quantity', 0)),
                'subtotal': f'¥{item_total:.2f}',
            })
        module_groups.append({
            'name': module.get('name', '未命名模块'),
            'items': items,
            'total': module_total
        })
    
    # 费用单独一组
    fees_total = sum(float(f.get('amount', 0)) for f in data.get('fees', []))
    fee_items = []
    for fee in data.get('fees', []):
        location_text = '厂内' if fee.get('location') == 'internal' else ('厂外' if fee.get('location') == 'external' else '')
        fee_items.append({
            'item': get_fee_type_name(fee.get('fee_type', ''), lang) or fee.get('name', ''),
            'spec': location_text,
            'category': get_fee_type_name(fee.get('fee_type', ''), lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f'¥{float(fee.get("amount", 0)):.2f}',
        })
    # 人力工时明细加入"其他费用"分组
    labor_total_sum = 0.0
    for lh in data.get('labor_hours', []):
        labor_total_sum += float(lh.get('total', 0))
        fee_items.append({
            'item': lh.get('name', t('labor_hours', lang)),
            'spec': '',
            'category': t('labor_hours', lang),
            'unit_price': f'¥{float(lh.get("unit_price", 0)):.2f}',
            'quantity': str(lh.get('hours', 0)),
            'subtotal': f'¥{float(lh.get("total", 0)):.2f}',
        })
    fees_total += labor_total_sum
    if fee_items:
        module_groups.append({
            'name': t('other_fees', lang),
            'items': fee_items,
            'total': fees_total
        })
    
    if module_groups:
        combined_table = doc.add_table(rows=1, cols=8)
        combined_table.style = 'Table Grid'
        headers = ['模块', '项目', '规格', '分类/品牌', '单价', '数量', '小计', '合计']
        for j, header in enumerate(headers):
            combined_table.rows[0].cells[j].text = header
        
        for group in module_groups:
            if not group['items']:
                row = combined_table.add_row()
                row.cells[0].text = group['name']
                row.cells[1].text = '（无物料）'
                continue
            
            first_row = combined_table.add_row()
            first_row.cells[0].text = group['name']
            first_row.cells[1].text = group['items'][0]['item']
            first_row.cells[2].text = group['items'][0]['spec']
            first_row.cells[3].text = group['items'][0]['category']
            first_row.cells[4].text = group['items'][0]['unit_price']
            first_row.cells[5].text = group['items'][0]['quantity']
            first_row.cells[6].text = group['items'][0]['subtotal']
            first_row.cells[7].text = f'¥{group["total"]:.2f}'
            
            if len(group['items']) > 1:
                module_cell = first_row.cells[0]
                for i in range(1, len(group['items'])):
                    next_row = combined_table.add_row()
                    next_row.cells[1].text = group['items'][i]['item']
                    next_row.cells[2].text = group['items'][i]['spec']
                    next_row.cells[3].text = group['items'][i]['category']
                    next_row.cells[4].text = group['items'][i]['unit_price']
                    next_row.cells[5].text = group['items'][i]['quantity']
                    next_row.cells[6].text = group['items'][i]['subtotal']
                    module_cell = module_cell.merge(next_row.cells[0])
                
                total_cell = first_row.cells[7]
                for i in range(1, len(group['items'])):
                    next_row_idx = len(combined_table.rows) - 1
                    total_cell = total_cell.merge(combined_table.rows[next_row_idx].cells[7])
    else:
        doc.add_paragraph('无物料及费用数据')
    
    doc.add_paragraph()
    
    # 汇总信息
    doc.add_heading(t('quote_summary', lang), level=1)
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Table Grid'
    fees_with_labor = totals['fees_total'] + totals['labor_total']
    summary_data = [
        (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
        (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
        (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
        (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
        (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
        (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
        (f'{t("grand_total", lang)}({currency})', f'{currency_symbol}{grand_total_converted:.2f}')
    ]
    for i, (key, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = value
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f'{data.get("name", "报价单")}_v{version_no}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@export_bp.route('/quotations/<int:quotation_id>/versions/<int:version_no>/export/pdf', methods=['GET'])
def export_version_pdf(quotation_id, version_no):
    lang = request.args.get('lang', 'zh')
    """导出版本 PDF 格式（直接下载预生成文件）"""
    version = VersionSnapshot.query.filter_by(
        quotation_id=quotation_id,
        version_no=version_no
    ).first()
    if not version:
        return jsonify({'error': 'Version not found'}), 404
    
    # 如果有预生成的文件，直接返回
    if version.pdf_file:
        import os
        import json
        # pdf_file 可能是 JSON 字符串 {'zh': path, 'en': path}
        pdf_paths = version.pdf_file
        if pdf_paths.startswith('{'):
            try:
                pdf_paths = json.loads(pdf_paths)
                file_path = pdf_paths.get(lang, pdf_paths.get('zh', ''))
            except:
                file_path = pdf_paths
        else:
            file_path = pdf_paths
        
        if file_path and os.path.exists(file_path):
            # snapshot_data 是 JSON 字符串
            name = '报价单'
            if version.snapshot_data:
                try:
                    data = json.loads(version.snapshot_data) if isinstance(version.snapshot_data, str) else version.snapshot_data
                    name = data.get('name', '报价单')
                except:
                    pass
            filename = f'{name}_v{version_no}_{lang}.pdf'
            return send_file(
                file_path,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=filename
            )
    
    # 如果没有预生成文件，动态生成
    data = get_version_snapshot_data(quotation_id, version_no)
    if not data:
        return jsonify({'error': 'Version data not found'}), 404
    
    # 获取报价单信息（包含币种）
    quotation = Quotation.query.get(quotation_id)
    
    totals = calculate_version_totals(data)
    
    # 使用报价单的币种
    currency = quotation.currency if quotation and quotation.currency else 'CNY'
    exchange_rates = {er.currency: float(er.rate) for er in ExchangeRate.query.all()}
    currency_symbol = get_currency_symbol(currency)
    grand_total_converted = convert_currency(totals['grand_total'], currency, exchange_rates)
    
    # 收集分组数据（乘系数）
    fee_rates = totals.get('fee_rates', {})
    rate_large = fee_rates.get('large', 1.0)
    rate_standard = fee_rates.get('standard', 1.0)
    rate_other = fee_rates.get('other', 1.0)

    module_groups = []
    for module in data.get('modules', []):
        items = []
        module_total = 0
        for mm in module.get('materials', []):
            cat = mm.get('category', 'standard')
            rate = rate_large if cat == 'large' else (rate_standard if cat == 'standard' else rate_other)
            unit_price_with_rate = float(mm.get('unit_price', 0) or 0) * rate
            item_total = unit_price_with_rate * float(mm.get('quantity', 0) or 0)
            module_total += item_total
            items.append({
                'item': mm.get('name', ''),
                'spec': mm.get('spec', ''),
                'category': mm.get('brand', ''),
                'unit_price': f'¥{unit_price_with_rate:.2f}',
                'quantity': str(mm.get('quantity', 0)),
                'subtotal': f'¥{item_total:.2f}',
            })
        module_groups.append({
            'name': module.get('name', t('unnamed_module', lang)),
            'items': items,
            'total': module_total
        })
    
    fees_total = sum(float(f.get('amount', 0)) for f in data.get('fees', []))
    fee_items = []
    for fee in data.get('fees', []):
        location_text = '厂内' if fee.get('location') == 'internal' else ('厂外' if fee.get('location') == 'external' else '')
        fee_items.append({
            'item': get_fee_type_name(fee.get('fee_type', ''), lang) or fee.get('name', ''),
            'spec': location_text,
            'category': get_fee_type_name(fee.get('fee_type', ''), lang) or '',
            'unit_price': '',
            'quantity': '1',
            'subtotal': f'¥{float(fee.get("amount", 0)):.2f}',
        })
    # 人力工时明细加入"其他费用"分组
    labor_total_sum = 0.0
    for lh in data.get('labor_hours', []):
        labor_total_sum += float(lh.get('total', 0))
        fee_items.append({
            'item': lh.get('name', '人力工时'),
            'spec': '',
            'category': '人力工时',
            'unit_price': f'¥{float(lh.get("unit_price", 0)):.2f}',
            'quantity': str(lh.get('hours', 0)),
            'subtotal': f'¥{float(lh.get("total", 0)):.2f}',
        })
    fees_total += labor_total_sum
    if fee_items:
        module_groups.append({
            'name': '其他费用',
            'items': fee_items,
            'total': fees_total
        })
    
    buffer = BytesIO()
    
    class PDF(FPDF):
        def header(self):
            pass
        def footer(self):
            pass
    
    pdf = PDF()
    pdf.add_font('SimHei', '', FONT_REGULAR, uni=True)
    pdf.add_font('SimHei', 'B', FONT_BOLD, uni=True)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # 标题
    pdf.set_font('SimHei', 'B', 16)
    pdf.cell(0, 12, data.get('name', '报价单'), 0, 1, 'C')
    pdf.ln(5)
    
    # 基本信息
    pdf.set_font('SimHei', 'B', 11)
    pdf.cell(0, 8, t('basic_info', lang), 0, 1)
    pdf.set_font('SimHei', '', 10)
    
    business_owner_name = ''
    if data.get('business_owner_id'):
        owner = User.query.get(data.get('business_owner_id'))
        if owner:
            business_owner_name = owner.real_name
    basic_info = [
        (t('project_name', lang), data.get('name', '')),
        (t('scheme_no', lang), data.get('scheme_no', '')),
        (t('business_owner', lang), business_owner_name),
        (t('version_no', lang), f'v{version_no}'),
        (t('currency', lang), currency)
    ]
    for key, value in basic_info:
        pdf.set_font('SimHei', 'B', 9)
        pdf.cell(40, 6, key, 1)
        pdf.set_font('SimHei', '', 9)
        pdf.cell(0, 6, str(value), 1, 1)
    
    pdf.ln(5)
    
    # 物料及费用清单 (PDF)
    if module_groups:
        pdf.set_font('SimHei', 'B', 11)
        pdf.cell(0, 8, t('material_list', lang), 0, 1)
        col_widths = [22, 38, 26, 20, 22, 12, 26, 24]  # 总约190
        headers = [t('module_col', lang), t('item_col', lang), t('spec_col', lang),
                   t('category_col', lang), t('unit_price_col', lang), t('qty_col', lang),
                   t('subtotal_col', lang), t('total_col', lang)]
        
        def draw_table_header():
            pdf.set_font('SimHei', 'B', 8)
            pdf.set_fill_color(68, 114, 196)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(pdf.l_margin if pdf.l_margin else 10)
            for i, h in enumerate(headers):
                pdf.cell(col_widths[i], 6, h, 1, 0, 'C', True)
            pdf.ln()
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('SimHei', '', 8)
        
        def draw_group(y_start, group):
            item_count = len(group['items'])
            row_h = 5
            y_bottom = y_start + item_count * row_h
            
            x_start = pdf.l_margin if pdf.l_margin else 10
            x_module = x_start
            x_data = x_start + col_widths[0]
            x_total = x_start + col_widths[0] + sum(col_widths[1:7])
            x_end = x_start + sum(col_widths)
            
            pdf.rect(x_module, y_start, col_widths[0], item_count * row_h)
            pdf.rect(x_total, y_start, col_widths[7], item_count * row_h)
            
            pdf.line(x_data, y_start, x_data, y_bottom)
            pdf.line(x_total, y_start, x_total, y_bottom)
            
            x_cur = x_data
            for j in range(1, 7):
                x_line = x_cur + col_widths[j]
                pdf.line(x_line, y_start, x_line, y_bottom)
                x_cur += col_widths[j]
            
            pdf.line(x_start, y_bottom, x_end, y_bottom)
            
            for i, item in enumerate(group['items']):
                y_item = y_start + i * row_h
                pdf.set_x(x_start)
                pdf.cell(col_widths[0], row_h, group['name'] if i == 0 else '', 0, 0, 'C')
                pdf.cell(col_widths[1], row_h, item['item'][:12], 0, 0, 'L')
                pdf.cell(col_widths[2], row_h, item['spec'][:8], 0, 0, 'L')
                pdf.cell(col_widths[3], row_h, item['category'][:6], 0, 0, 'L')
                pdf.cell(col_widths[4], row_h, item['unit_price'][:7], 0, 0, 'R')
                pdf.cell(col_widths[5], row_h, item['quantity'], 0, 0, 'C')
                pdf.cell(col_widths[6], row_h, item['subtotal'][:8], 0, 0, 'R')
                pdf.cell(col_widths[7], row_h, f'¥{group["total"]:.2f}' if i == 0 else '', 0, 1, 'R')
            
            return y_bottom
        
        draw_table_header()
        y = pdf.get_y()
        for group in module_groups:
            if group['items']:
                y = draw_group(y, group)
        
        pdf.ln(5)
        
        # 报价汇总
        pdf.set_font('SimHei', 'B', 11)
        pdf.cell(0, 8, t('quote_summary', lang), 0, 1)
        
        pdf.set_font('SimHei', 'B', 9)
        pdf.set_fill_color(68, 114, 196)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(100, 6, t('project_col', lang), 1, 0, 'C', True)
        pdf.cell(0, 6, t('amount_col', lang), 1, 1, 'C', True)
        
        pdf.set_text_color(0, 0, 0)
        fees_with_labor = totals['fees_total'] + totals['labor_total']
        summary_items = [
            (t('material_total', lang), f'¥{totals["material_total_with_rates"]:.2f}'),
            (t('fee_total', lang), f'¥{fees_with_labor:.2f}'),
            (t('profit_rate', lang), f'{totals["profit_rate"] * 100:.1f}%'),
            (t('subtotal_with_profit', lang), f'¥{totals["subtotal_with_profit"]:.2f}'),
            (t('tax_rate', lang), f'{totals["tax_rate"] * 100:.0f}%'),
            (t('tax_amount', lang), f'¥{totals["tax_amount"]:.2f}'),
        ]
        pdf.set_font('SimHei', '', 9)
        for label, value in summary_items:
            pdf.cell(100, 5, label, 1)
            pdf.cell(0, 5, value, 1, 1, 'R')
        
        # 最终报价（高亮）
        pdf.set_font('SimHei', 'B', 10)
        pdf.set_fill_color(226, 239, 218)
        pdf.cell(100, 7, f"{t('grand_total', lang)}({currency})", 1, 0, 'C', True)
        pdf.cell(0, 7, f'{currency_symbol}{grand_total_converted:.2f}', 1, 1, 'R', True)
    
    pdf_bytes = pdf.output()
    buffer = BytesIO(pdf_bytes)
    filename = f'{data.get("name", t("quote", lang))}_v{version_no}_{datetime.now().strftime("%Y%m%d%H%M%S")}.pdf'
    return send_file(buffer, mimetype='application/pdf', as_attachment=True, download_name=filename)
